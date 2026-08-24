"""
DOCX质量回归测试
每次修改导出管线后运行，确保不引入已知问题
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from docx import Document
from docx.oxml.ns import qn

SAMPLE_DOCX = os.path.join(os.path.dirname(__file__), "..", "output", "芯联集成_cicc_clean.docx")


@pytest.mark.skipif(not os.path.exists(SAMPLE_DOCX), reason="需要先有报告")
class TestDocxQuality:
    
    def test_no_bold_markers(self):
        """DOCX中不应该有 ** 标记"""
        doc = Document(SAMPLE_DOCX)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "**" not in text, "发现 ** 标记残留"
    
    def test_no_table_separators(self):
        """DOCX中不应该有 :--- 表格分隔符"""
        doc = Document(SAMPLE_DOCX)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert ":---" not in text, "发现 :--- 表格分隔符残留"
    
    def test_first_element_not_empty(self):
        """文档首元素不能是空表或空段落"""
        doc = Document(SAMPLE_DOCX)
        body = doc.element.body
        assert len(body) > 0, "文档体为空"
        first = body[0]
        tag = first.tag.split("}")[-1] if "}" in first.tag else first.tag
        if tag == "tbl":
            cells = first.findall(".//" + qn("w:t"))
            assert not all(c.text is None or c.text.strip() == "" for c in cells), "首元素为空表格"
        elif tag == "p":
            texts = [t.text for t in first.findall(".//" + qn("w:t")) if t.text]
            assert not all(t.strip() == "" for t in texts), "首段落为空"
    
    def test_has_images(self):
        """报告至少要有图片"""
        doc = Document(SAMPLE_DOCX)
        assert len(doc.inline_shapes) > 0, "没有内嵌图片"
    
    def test_has_tables(self):
        """报告至少要有表格"""
        doc = Document(SAMPLE_DOCX)
        assert len(doc.tables) > 0, "没有表格"
    
    def test_first_para_is_title(self):
        """首段落应该是标题"""
        doc = Document(SAMPLE_DOCX)
        if doc.paragraphs:
            assert len(doc.paragraphs[0].text.strip()) > 0, "首段落为空"
    
    def test_visual_gate_passes(self):
        """VisualGate 结构校验：回归验证门禁拦截能力。

        R6（2026-08-01 圆桌升级）：VisualGate 从白名单否决升级为结构校验。
        此前白名单版只查"有没有坏模式"，对"没有结构/图表堆末尾/来源附录重复"
        这类缺失型缺陷完全免疫（章节缺失→正则空匹配反而通过，图表全堆末尾→图片数达标算通过）。
        本测试构造自包含的缺陷样本验证新检查生效，不依赖 output 目录运行产物。

        关键：真实修复后的报告（传感器行业_cicc.docx）应能通过门禁；
        但门禁必须能拦截"无章节/图表堆末尾/孤立句号"的坏样本。
        """
        from docx import Document as _Doc
        import tempfile
        from export.visual_gate import check

        def _make_doc(paras_with_style, inline_shapes=0):
            """构造 DOCX：[(text, style_name)] 列表"""
            d = _Doc()
            for txt, style in paras_with_style:
                p = d.add_paragraph(txt, style=style)
            return d

        def _save(d, name):
            path = os.path.join(tempfile.mkdtemp(), name)
            d.save(path)
            return path

        # —— 坏样本1：无章节结构（正文全是 Normal 平铺，无 Heading 文本）——
        bad1 = _make_doc([
            ("给予传感器行业增持评级。市场普遍认为……", "Normal"),
            ("。全球传感器市场规模从2023年增长。", "Normal"),
            ("产业链上游中游下游分析。", "Normal"),
        ])
        bad1_path = _save(bad1, "bad1_no_structure.docx")
        r1 = check(bad1_path, "industry_deep")
        assert not r1["passed"], "无章节结构样本应被拦截"
        assert "missing_section_structure" in {i["check"] for i in r1["issues"]}, \
            "missing_section_structure 未生效: %s" % [i["check"] for i in r1["issues"]]

        # —— 坏样本2：孤立段首句号 + 来源附录重复 ——
        bad2 = _make_doc([
            ("报告标题", "Heading 1"),
            ("一、市场空间", "Heading 2"),
            ("。市场从2023年增长。", "Normal"),  # 孤立句号
            ("二、竞争格局", "Heading 2"),
            ("。竞争加剧。", "Normal"),
            ("三、利润池", "Heading 2"),
            ("利润迁移。", "Normal"),
            ("<!-- AGENT_ENRICH_SOURCES -->来源附录", "Normal"),
            ("<!-- AGENT_ENRICH_SOURCES -->来源附录重复", "Normal"),
        ])
        bad2_path = _save(bad2, "bad2_layout.docx")
        r2 = check(bad2_path, "industry_deep")
        assert not r2["passed"], "孤立句号+来源附录重复样本应被拦截"
        checks2 = {i["check"] for i in r2["issues"]}
        assert "stray_leading_period" in checks2 or "duplicate_source_appendix" in checks2, \
            "排版卫生检查未生效: %s" % [i["check"] for i in r2["issues"]]

        # —— 好样本：结构完整（≥3 标题），应通过结构校验 ——
        good = _make_doc([
            ("报告标题", "Heading 1"),
            ("一、市场空间", "Heading 2"),
            ("市场从2023年1797亿美元增长。", "Normal"),
            ("二、竞争格局", "Heading 2"),
            ("竞争加剧，龙头集中。", "Normal"),
            ("三、利润池", "Heading 2"),
            ("利润向IDM迁移。", "Normal"),
        ])
        good_path = _save(good, "good.docx")
        r3 = check(good_path, "industry_deep")
        # 结构校验应通过（但可能因字体/表格等白名单项失败，这里只验证结构检查不再报错）
        assert "missing_section_structure" not in {i["check"] for i in r3["issues"]}, \
            "结构完整样本不应报 missing_section_structure"

        # —— 已知缺陷旧样本（若存在）：应被拦截 ——
        if os.path.exists(SAMPLE_DOCX):
            r_old = check(SAMPLE_DOCX, "listed_company")
            assert not r_old["passed"], "已知缺陷样本应被 VisualGate 拦截，但通过了"
