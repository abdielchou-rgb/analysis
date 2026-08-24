# -*- coding: utf-8 -*-
"""R28 事实质量全量修复测试（数据口径 + Gate一致性 + 写作规划）

固化四刀修复的关键行为：
  1. 数据冲突检测（毛利率/PE/营收多来源矛盾 → 检测）
  2. Gate 评级-空间一致性（+2.7%给增持 → 拦截）
  3. Gate 估值锚一致性（PE法 vs DCF法 >20% → 拦截）
  4. 写作规划生成（必答问题 + 结论自洽约束）
"""

import os
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))


# ── 测试 1：数据冲突检测 ────────────────────────────────────────
def test_data_conflict_detection():
    from core.data_caliber import detect_value_conflicts

    # 柯力场景：毛利率 34.5%（全年）vs 5.0%（单季）
    data = {"fig_margin_2025": 34.5, "fig_margin_2026q1": 5.0}
    conflicts = detect_value_conflicts(data)
    margin_conflicts = [c for c in conflicts if c["indicator"] == "margin"]
    assert margin_conflicts, "毛利率多来源矛盾应被检测"
    assert margin_conflicts[0]["severity"] == "error"
    assert margin_conflicts[0]["gap_pct"] > 100

    # PE 冲突（65 vs 79.79）
    data2 = {"pe_ttm": 65, "fig_valuation_pe": 79.79}
    conflicts2 = detect_value_conflicts(data2)
    pe_conflicts = [c for c in conflicts2 if c["indicator"] == "pe"]
    assert pe_conflicts, "PE 多来源矛盾应被检测"

    # 无冲突
    data3 = {"revenue_2025": 15.58}
    assert detect_value_conflicts(data3) == []


# ── 测试 2：口径标注 ────────────────────────────────────────────
def test_caliber_annotation():
    from core.data_caliber import build_caliber_meta, serialize_caliber_annotations

    meta = build_caliber_meta({"revenue_2025": 15.58, "pe_ttm": 65})
    assert meta["revenue_2025"]["unit"] == "亿元", "营收应标注亿元"
    assert meta["pe_ttm"]["unit"] == "倍", "PE应标注倍"
    s = serialize_caliber_annotations(meta)
    assert "unit=" in s, "口径标注应含单位"


# ── 测试 3：Gate 评级-空间一致性 ────────────────────────────────
def test_gate_rating_target_consistency():
    from pipeline.iron_gate import IronGate
    import tempfile

    # 柯力场景：增持但仅+2.7%空间
    report = (
        "# 测试报告\n投资评级：增持，12个月目标价48元，现价46.73元。\n"
        "公司2025年营收15.58亿元，毛利率34.5%。DCF公允市值约150亿元。\n"
        "继续补充内容触发检查逻辑，保证正文超过三百字阈值。"
        "北向资金净流出，两融余额较低，筹码分散。商业模式为能者模型。"
        "护城河中等，转换成本中等，品牌护城河弱。财务政策审慎。"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_rating_target_consistency()
    os.unlink(tmp.name)
    assert not r.passed, "增持+2.7%空间应被拦截"
    assert "评级-空间错配" in r.details, "应指出评级-空间错配"


# ── 测试 4：Gate 数据冲突检查 ──────────────────────────────────
def test_gate_data_conflicts_check():
    from pipeline.iron_gate import IronGate
    import tempfile

    report = "# 测试\n毛利率34.5%，毛利率5.0%，营收15.58亿元。" * 5
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.asset = "603662"
    r = ig._check_data_conflicts()
    os.unlink(tmp.name)
    # 无 data_dict 缓存时降级 warning（不阻断），但方法可运行
    assert r.name == "data_conflicts"


# ── R32 回归：Gate 目标价自相矛盾检测（柯力案 51.60 vs 48）──────
def test_gate_multiple_target_prices_conflict():
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "**投资评级：增持** ｜ **12个月目标价：51.60元** ｜ **当前价：46.73元**\n"
        "公司2025年营收15.58亿元，毛利率34.5%，净利率10.8%。"
        "DCF公允市值约145-160亿元。PE估值对应目标价40-48元。"
        "综合DCF+PE，12个月目标价48元，给予增持评级。"
        "北向资金增持，两融余额较低，筹码集中。护城河中等。"
        "商业模式为能者模型。财务政策审慎。行业景气中性偏弱。"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_rating_target_consistency()
    os.unlink(tmp.name)
    assert not r.passed, "两个矛盾目标价应被拦截"
    assert "目标价自相矛盾" in r.details, "应指出目标价矛盾"


def test_gate_single_target_price_passes():
    """R32：单一目标价 + 区间表述不应误报。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "**投资评级：增持** ｜ **12个月目标价：48.00元** ｜ **当前价：43.00元**\n"
        "PE估值合理区间40-48元，DCF估值约50-55元，综合取48元作为目标价。"
        "公司2025年营收15.58亿元，毛利率34.5%。北向资金增持。"
        "护城河中等，转换成本中等。商业模式为能者模型。财务政策审慎。"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_rating_target_consistency()
    os.unlink(tmp.name)
    assert r.passed, "单一目标价不应误报"


# ── R34 回归：so_what_chain 表格段豁免 ─────────────────────────
def test_so_what_table_sections_exempt():
    """R34：纯表格段（跟踪指标表/风险表）不应拉低 so_what min_score。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = """# 测试报告
## 一、行业分析

2025年全球市场规模890亿美元，因此我们判断行业处于成长期。数据表明渗透率持续提升，这意味着国产替代空间大。

## 二、竞争格局

公司市占率第一，因此我们认为护城河稳固。反方论证：低价竞争可能侵蚀份额，我们判断需跟踪价格趋势。

## 三、跟踪指标

| 维度 | 指标 | 当前值 | 验证信号 |
|------|------|--------|---------|
| 增长 | 营收增速 | 20% | 连续两季超25% |
| 盈利 | 毛利率 | 34.5% | 站稳42%以上 |
| 资金 | 北向持股 | 318万股 | 持续增持 |

## 四、风险提示

| 风险 | 概率 | 影响 |
|------|------|------|
| 估值回调 | 高 | 高 |
| 商誉减值 | 低 | 中 |
"""
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_so_what_chain()
    os.unlink(tmp.name)
    assert r.passed, f"表格段应被豁免，当前 min_score 不应被拉低: {r.details}"


# ── R35 回归：算术校验层（占比/估值中值/目标价空间/EPS桥）──
def test_arithmetic_audit_catches_wrong_ratio():
    """R35：北向占比算错（0.24% vs 实际1.13%）应被算术校验拦截。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "**投资评级：增持** ｜ **12个月目标价：53.50元** ｜ **当前价：46.73元** ｜ **目标价空间：+14.5%**\n"
        "## 资金面\n"
        "北向资金持股318.29万股，占总股本约0.24%（基于总市值131.23亿元、股价46.73元推算总股本约2.81亿股）。\n"
        "融资余额5.70亿元，占流通市值比4.17%。股东户数环比下降11.97%。\n"
        "## 估值\n"
        "目标价区间为42-64元，中值约53.50元与正文目标价一致。2027E EPS 1.07元。\n"
        "## 风险\n"
        "行业价格战风险，机器人放量不及预期风险。公司毛利率结构性上行。\n"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_arithmetic_audit()
    os.unlink(tmp.name)
    assert not r.passed, "北向占比算错应被拦截"
    assert "占比验算错误" in r.details, "应指出占比验算错误"


def test_arithmetic_audit_passes_correct_ratio():
    """R35：正确的占比不应误报。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "**投资评级：增持** ｜ **12个月目标价：53.50元** ｜ **当前价：46.73元** ｜ **目标价空间：+14.5%**\n"
        "## 资金面\n"
        "北向资金持股318.29万股，占总股本约1.13%。融资余额5.70亿元，占流通市值比4.17%。\n"
        "## 估值\n"
        "目标价区间为42-64元，中值约53.0元。2027E EPS 1.07元。\n"
        "## 风险\n"
        "行业价格战风险。公司毛利率结构性上行，财务健康。\n"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_arithmetic_audit()
    os.unlink(tmp.name)
    assert r.passed, f"正确占比不应误报: {r.details}"


# ── R35 回归：模板句高重复检测 ─────────────────────────────
def test_template_repeat_catches_pollution():
    """R35：模板句重复 2 次 + 概念错位应被检测。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    base = (
        "## 分析\n"
        "公司主营称重传感器，2025年营收15.58亿元，同比增长20.33%。因此，这意味着公司增长强劲。"
        "这一趋势若持续，盈利中枢存在系统性上移的可能。\n"
        "## 估值\n"
        "端侧变现的兑现节奏是关键变量。这一趋势若持续，盈利中枢存在系统性上移的可能。\n"
    )
    # 确保超过 300 字阈值
    filler = "公司毛利率结构性上行，财务健康，护城河中等偏弱。DCF与PE交叉验证，目标价自洽。" * 4
    report = base + base + filler
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_template_repeat()
    os.unlink(tmp.name)
    assert not r.passed, "模板句重复应被检测"
    assert "模板句重复" in r.details, "应指出模板句重复"


def test_template_repeat_passes_clean():
    """R35：正常报告不应误报模板污染。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "## 分析\n"
        "公司2025年营收15.58亿元，同比增长20.33%，因此我们判断增长质量是核心。\n"
        "## 估值\n"
        "DCF与PE两法交叉验证，目标价自洽。\n"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    r = ig._check_template_repeat()
    os.unlink(tmp.name)
    assert r.passed, f"正常报告不应误报: {r.details}"


# ── R38 回归：财务数值一致性（毛利率/PE 与 data_dict 冲突）──
# 注意：测试用独立临时 asset 名（_test_keli），不触碰真实 output/柯力传感_data_dict.json，
# 避免污染真实管线数据（曾发生过测试覆盖真实 data_dict 的事故）。
_TEST_DD_PATH = Path("output/_test_keli_data_dict.json")


def test_financial_value_consistency_catches_conflict():
    """R38：报告毛利率与 data_dict 真实值冲突应被拦截（柯力案 34.5% vs 46.35%）。"""
    from pipeline.iron_gate import IronGate
    import tempfile, json
    from pathlib import Path

    # 写入临时 data_dict（独立 asset 名，不污染真实柯力数据）
    dd = {"margin_2025": 44.83, "margin_2026": 46.35}
    Path(_TEST_DD_PATH).write_text(json.dumps(dd, ensure_ascii=False), encoding="utf-8")

    report = (
        "## 财务分析\n"
        "公司毛利率自2018年40.87%回落至2025年34.5%附近。当前毛利率维持34.5%左右。\n"
        "营收15.58亿元，净利1.68亿元。\n"
        "## 估值\n"
        "动态PE 79.79倍，附录PE(TTM) 44.63倍。\n"
        "## 风险\n"
        "行业价格战风险，机器人放量不及预期。公司毛利率结构性上行，财务健康。\n"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.asset = "_test_keli"
    r = ig._check_financial_value_consistency()
    os.unlink(tmp.name)
    assert not r.passed, "毛利率矛盾应被拦截"
    assert "毛利率矛盾" in r.details, "应指出毛利率矛盾"


def test_financial_value_consistency_passes_good():
    """R38：与 data_dict 一致的财务值不应误报。"""
    from pipeline.iron_gate import IronGate
    import tempfile, json
    from pathlib import Path

    dd = {"margin_2025": 44.83, "margin_2026": 46.35}
    Path(_TEST_DD_PATH).write_text(json.dumps(dd, ensure_ascii=False), encoding="utf-8")

    report = (
        "## 财务分析\n"
        "公司2025年毛利率约44.8%，维持高位。营收15.58亿元，净利1.68亿元。\n"
        "## 估值\n"
        "动态PE 79.79倍（2025静态口径），2027E前瞻PE约43.7倍。\n"
        "## 风险\n"
        "行业价格战风险。公司毛利率结构性上行，财务健康。\n"
    ) * 2
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.asset = "_test_keli"
    r = ig._check_financial_value_consistency()
    os.unlink(tmp.name)
    assert r.passed, f"一致财务值不应误报: {r.details}"


# ── R38 回归：导出残留清理（_scrub_aigc_artifacts）──
def test_scrub_aigc_artifacts():
    """R38：AGENT_ENRICH 注释/内部字段名/病句/括号计数应从成品清除。"""
    from export.docx_exporter import _scrub_aigc_artifacts

    test = (
        "## 分析\n"
        "公司2025年营收15.58亿元。。毛利率44.8%。，净利率10.8%。\n"
        "公司研究素材.geopolitical_summary 提到地缘风险。\n"
        "共6个环节（6）驱动因素计数为5项（5）\n"
        "<!-- AGENT_ENRICH_SOURCES -->\n"
        "来源标注\n"
        "<!-- /AGENT_ENRICH_SOURCES -->\n"
        "正常正文内容。\n"
    )
    clean = _scrub_aigc_artifacts(test)
    assert "AGENT_ENRICH" not in clean, "HTML 注释应清除"
    assert "公司研究素材" not in clean, "内部字段名应清除"
    assert "。。" not in clean, "病句应清除"
    assert "。，" not in clean, "病句应清除"
    assert "正常正文内容" in clean, "正常内容应保留"


# ── R40 回归：渲染层目检（docx 空段/分页/图表分布）──
def test_layout_quality_catches_docx_issues():
    """R40：渲染层目检应捕获 docx 空段落率过高/连续空段/图表未随文。"""
    from pipeline.iron_gate import IronGate
    import tempfile, zipfile, shutil
    from pathlib import Path

    # 构造一个含空段 + 图表集中的 docx
    import docx
    doc = docx.Document()
    # 15 个空段
    for _ in range(15):
        doc.add_paragraph("")
    doc.add_paragraph("## 财务分析")
    doc.add_paragraph("公司2025年营收15.58亿元，毛利率44.8%。" * 5)
    # 图表集中在最后
    for _ in range(3):
        doc.add_paragraph("![图](charts/fig1.png)")
    doc.save("output/_layout_test.docx")

    # 构造 Gate 检查
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write("## 财务分析\n公司2025年营收15.58亿元，毛利率44.8%。\n" * 10)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = open(tmp.name, encoding="utf-8").read()
    # 让 layout_quality 找到测试 docx
    ig.report_path = "output/_layout_test.md"
    Path("output/_layout_test.md").write_text(ig.report_text, encoding="utf-8")
    r = ig._check_layout_quality()
    os.unlink(tmp.name)
    os.unlink("output/_layout_test.md")
    os.unlink("output/_layout_test.docx")
    # 空段率 >15% 应触发
    assert not r.passed, "空段率过高应被检测"
    assert "空段落率" in r.details, "应指出空段落率"


# ── R46 回归：不变量断言层（物理不可能拦截）────────────────
def test_invariant_audit_catches_holding_value():
    """R46：持股数×股价≠持股市值（r11 案 318.29万股×46.73=1.49亿 vs 2.27亿）应拦截。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "总市值131.23亿元，收盘价46.73元。北向资金持有柯力传感318.29万股，持股市值2.27亿元。\n"
        "融资余额5.70亿元，占流通市值比为4.17%。PE(TTM) 78.1倍，净利1.68亿元。\n"
        "正文内容充足，财务数据自洽。\n"
    ) * 8
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.report_type = "listed_company"
    r = ig._check_invariant_audit()
    os.unlink(tmp.name)
    assert not r.passed, "持股市值矛盾应被拦截"
    assert "持股市值矛盾" in r.details, "应指出持股市值矛盾"


def test_invariant_audit_passes_good():
    """R46：自洽数据不应误报。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "总市值131.23亿元。北向资金持有柯力传感318.29万股，持股市值1.49亿元。\n"
        "融资余额5.70亿元，占流通市值比为4.17%。PE(TTM) 78.1倍，净利1.68亿元。\n"
        "正文内容充足，财务数据自洽。\n"
    ) * 8
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.report_type = "listed_company"
    r = ig._check_invariant_audit()
    os.unlink(tmp.name)
    assert r.passed, f"自洽数据不应误报: {r.details}"


def test_invariant_audit_catches_dcf_circular():
    """R46：DCF 循环论证（r11 案 模型39亿 vs 报告145-160亿）应拦截。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    report = (
        "DCF公允市值区间为145-160亿元。FCFF 约 1.75 亿元，WACC 8.5%，永续增长率2.5%。"
        "2026E +15%、2027E +12%、2028E +10%。\n"
        "正文内容充足。\n"
    ) * 8
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(report)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = report
    ig.report_type = "listed_company"
    r = ig._check_invariant_audit()
    os.unlink(tmp.name)
    assert not r.passed, "DCF 循环论证应被拦截"
    assert "DCF循环论证" in r.details, "应指出 DCF 循环论证"


# ── R41 回归：frontmatter 水印豁免 + 免责声明白名单 ──────────
def test_md_artifacts_frontmatter_exempt():
    """R41：AIGC 水印 frontmatter（--- YAML 头）不应被误报为多余分隔符。"""
    from pipeline.iron_gate import IronGate
    import tempfile

    text = (
        "---\n"
        "AIGC:\n"
        "    Label: \"1\"\n"
        "    ContentProducer: xxx\n"
        "---\n"
        "# 柯力传感深度报告\n"
        "公司2025年营收15.58亿元，毛利率44.8%。正文内容充足无残留。\n"
        "## 估值\n"
        "目标价53.50元，评级增持。\n"
    )
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
    tmp.write(text)
    tmp.close()
    ig = IronGate(tmp.name, "listed_company", "cicc")
    ig.report_text = text
    r = ig._check_markdown_artifacts()
    os.unlink(tmp.name)
    assert r.passed, f"frontmatter 不应误报: {r.details}"


def test_forbidden_patterns_disclaimer_hardkill():
    """R72（2026-08-05）：AI 免责声明已从豁免改为硬拦截。

    R42 要求清除所有 AI 免责声明；但 R72 圆桌审计发现 Marvis 手动修复路径
    让"内容由AI生成，仅供参考"在手动路径复活。现改为硬拦截——该模式出现即 FAIL，
    任何路径（含手动修复）写出的报告都不允许携带 AI 免责声明。
    """
    from pipeline.iron_gate import IronGate

    text = (
        "# 柯力传感深度报告\n"
        "公司2025年营收15.58亿元，毛利率44.8%。正文内容充足。\n"
        "*（内容由AI生成，仅供参考）*\n"
    )
    ig = IronGate.__new__(IronGate)
    ig.report_text = text
    r = ig._check_forbidden_patterns()
    assert not r.passed, f"AI免责声明应被硬拦截(FAIL): {r.details}"
    assert "P0" in r.details, f"应标记P0级问题: {r.details}"

    # 对照组：无免责声明的正文应通过
    ig2 = IronGate.__new__(IronGate)
    ig2.report_text = "# 柯力传感深度报告\n公司2025年营收15.58亿元，毛利率44.8%。正文内容充足。\n"
    r2 = ig2._check_forbidden_patterns()
    assert r2.passed, f"正常正文不应误报: {r2.details}"


# ── R42 回归：不注入 AI 免责声明 + DOCX 静态目录 ──────────────
def test_no_ai_disclaimer_injection():
    """R42：professionalize 不应注入 AI 免责声明（报告须像人类分析师撰写）。"""
    from export.format_professionalizer import FormatProfessionalizer
    fp = FormatProfessionalizer()
    text = "# 柯力传感深度报告\n\n公司2025年营收15.58亿元，毛利率44.8%。\n"
    out = fp.professionalize(text)
    assert "免责声明" not in out, "不应注入免责声明"
    assert "仅供参考" not in out, "不应注入仅供参考"
    assert "不构成投资建议" not in out, "不应注入投资建议"
    assert "数据来源" in out, "数据来源是专业报告的合理部分，应保留"


def test_docx_static_toc_inserted():
    """R42：docx 导出应插入静态目录（含章节标题，无需 Word 刷新）。"""
    from export.docx_exporter import markdown_to_docx, add_static_toc
    import os

    md = (
        "# 柯力传感深度报告\n"
        "# 一、核心投资判断\n正文内容。\n"
        "## 二、财务分析\n正文内容。\n"
        "### 2.1 营收\n详细数据。\n"
        "## 三、估值\n正文内容。\n"
    )
    out = "output/_toc_regression.docx"
    markdown_to_docx(md, out, title="柯力传感深度报告")
    n = add_static_toc(out, md)
    # R43：应包含 1 个一级（# 一、）+ 2 个二级（## 二、/三、）+ 1 个三级（### 2.1）= 4
    assert n == 4, f"应插入4个章节标题（含一级章节）: {n}"
    from docx import Document
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert "目  录" in joined, "应有目录标题"
    assert "一、核心投资判断" in joined, "目录应含一级章节"
    assert "2.1 营收" in joined, "目录应含三级章节"
    os.unlink(out)





# ── 测试 5：写作规划生成 ────────────────────────────────────────
def test_report_plan():
    from core.report_planner import build_report_plan, serialize_plan

    plan = build_report_plan("listed_company")
    assert plan["total_critical"] >= 5, "上市公司应有至少5个必答问题"
    assert "评级" in " ".join(plan["must_answer"]), "必答问题应含评级"
    s = serialize_plan(plan)
    assert "必答" in s and "自洽" in s
    assert "≥10%" in s, "自洽约束应含评级空间要求"

    # 三种类型都可用
    for rt in ["listed_company", "industry_deep", "unlisted_company"]:
        p = build_report_plan(rt)
        assert len(p["questions"]) >= 5, f"{rt} 至少5个问题"


if __name__ == "__main__":
    import traceback

    passed = 0
    failed = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn()
                print(f"  ✓ {name}")
                passed += 1
            except Exception as e:
                print(f"  ✗ {name}: {e}")
                traceback.print_exc()
                failed += 1
    print(f"\n{passed} passed, {failed} failed")
    sys.exit(1 if failed else 0)
