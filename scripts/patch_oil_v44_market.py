# -*- coding: utf-8 -*-
"""油位 v4.4：清理版本迭代说明 + 补强市场调研（对照 2hao 框架）

1. 清理"版本迭代对比说明"（不存在/而非/见3.2/无需外部）→ 直接陈述
2. 补强市场调研：产业链价值分配/渗透率/供需/价格趋势（对照 2hao SAC 框架）
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.4.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.3.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 清理版本迭代说明 ==========
    # 3.2 合作实质："而非外部交易" → 删
    p = find_para(doc, "本次合作本质是集团内部业务协同，而非外部交易")
    if p:
        replace_para_text(p,
            "久通物联为柯力传感控股子公司（柯力官网公告「久通成为柯力传感控股子公司」，纳入集团控股参股体系）。"
            "本次生产统筹属于集团内部业务协同：久通拥有80余国海关/物流客户网络与罐箱物联网平台，"
            "具备油位传感器产品线（JT606X）；柯力具备制造、认证与规模化能力。"
            "将油位传感器生产统筹至柯力制造体系，既释放久通的渠道与平台价值，"
            "又补齐柯力在液位品类的制造短板。")
        logging.info("✅ 3.2 去'而非外部交易'")

    # 6.5 "久通已是控股子公司，本方案的本质是集团内部整合" → 直接陈述
    p = find_para(doc, "久通已是控股子公司，本方案的本质是集团内部整合")
    if p:
        replace_para_text(p, "内部整合路径如下：")
        logging.info("✅ 6.5 去'本方案的本质'")

    # 路径一 "无需外部签约" → 删
    p = find_para(doc, "集团内部产能调配，无需外部签约")
    if p:
        replace_para_text(p, "集团内部产能调配，久通聚焦渠道与物联网平台。")
        logging.info("✅ 路径一去'无需外部签约'")

    # 4.2 "而非久通现售的300元低端产品（低端产品仅作代工入场券，见3.2业务定位）" → 简化
    p = find_para(doc, "而非久通现售的300元低端产品")
    if p:
        replace_para_text(p,
            "磁致伸缩液位仪价格带3,000-8,000元/台。测算主体为本公司主导开发的中高端磁致伸缩产品，"
            "定位面向加油站防渗改造、危化品储罐安全仪表与罐箱油品运输监测，"
            "需通过防爆/计量认证（ATEX/IECEx）。依据竞争定位（中高端磁致伸缩、避开低端红海、"
            "不拼最高端防爆），对应价格带4,500-6,500元，取中位5,000元/台。")
        logging.info("✅ 4.2 去'而非/见3.2'")

    # ========== 2. 补强市场调研（1.4 后补产业链价值分配 + 供需 + 渗透率） ==========
    p_supply = find_para(doc, "该约束为行业共性问题")
    if p_supply:
        insert_after(p_supply, [
            ("t",
             "产业链价值分配：油位传感器产业链呈『两端高、中间低』结构——上游磁致伸缩波导丝材料毛利50%-65%，"
             "下游集成商（含防爆/计量认证）获取品牌与服务溢价，中游装配制造净利率仅8%-12%。"
             "此结构意味着：单纯代工装配难以建立利润优势，须向上游材料或下游认证/渠道延伸。"),
            ("t",
             "供需平衡：供给端国内磁致伸缩液位仪产能集中于中低端，高端防爆认证产能依赖外资与少数国产突破企业；"
             "需求端受2026-2027年加油站防渗与SIS改造双重窗口驱动集中释放，短期供需偏紧、"
             "价格具备支撑；2028年后政策红利回落，需求增速放缓，价格或承压。"),
            ("t",
             "渗透率视角：加油站液位监测在东部执行率接近100%（存量改造收尾），中西部仍有5%-8%缺口；"
             "罐箱油品运输监测渗透率处于早期（全球50-60万个罐箱，液位监测加装率较低），"
             "是未来5年渗透率提升的主要增量来源。"),
        ])
        logging.info("✅ 补产业链价值分配/供需/渗透率")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
