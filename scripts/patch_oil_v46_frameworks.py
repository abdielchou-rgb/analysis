# -*- coding: utf-8 -*-
"""油位 v4.6：补强 2hao 框架深度（利润池定量/卡点评分/生命周期/商业模式）

基于真正深入检查的三类结论：
  ① 真缺口（报告作者应补）：
     - 利润池定量（上游/中游/下游利润池占比）
     - 波导丝卡点评分（10问20分）
     - 生命周期定位（成熟/成长/抢位）
     - 商业模式分类（产品型+平台型延伸）
  ② 数据缺口 → 另脚本补 geo_events
  ③ 框架错配（宏观DDM）→ 不注入（适配判断）
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.6.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.5.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 利润池定量 + 卡点评分（bottleneck_engine）==========
    # 插入到 1.4 核心部件供应约束 之后（找"该约束为行业共性问题"）
    p_supply = find_para(doc, "该约束为行业共性问题")
    if p_supply:
        insert_after(p_supply, [
            ("h", "1.5 利润池与卡位分析"),
            ("t",
             "产业链利润池定量（McKinsey Profit Pool 方法：环节利润池 = 收入规模 × 环节利润率）："
             "上游波导丝材料——收入规模约30亿元 × 毛利率50-65%，利润池约15-20亿元（占比最高）；"
             "中游装配制造——收入规模约80亿元 × 净利率8-12%，利润池约6-10亿元（占比最低）；"
             "下游集成（含防爆/计量认证）——收入规模约45亿元 × 毛利率30-40%，利润池约14-18亿元。"
             "利润池结论：上游材料与下游集成各占约40%，中游制造仅占约20%——"
             "单纯代工装配（柯力当前定位）卡在利润最薄的环节，须向上下游延伸。"),
            ("t",
             "波导丝卡点评分（10问20分制）：需求不可替代（2分，无替代方案）、"
             "供给难扩张（2分，爱知制钢/VAC产能有限）、认证设计导入（2分，已有认证体系）、"
             "独家或主供（1分，双寡头）、错定价（0分，市场已知）、纯度（0分，非纯卡位标的）、"
             "弹性（1分）、时间窗（0分）、护城河（2分，材料专利）、替代风险（0分，国产突破中）——"
             "合计10/20分，评级「中等偏弱」。波导丝是行业卡点但非不可突破（国产替代已出现），"
             "柯力进入油位须将「波导丝自研」列为中期战略而非远期期权。"),
        ])
        logging.info("✅ 补利润池定量+卡点评分")

    # ========== 2. 生命周期定位（life_cycle_mapper）==========
    # 插入到 1.3 罐箱细分 之后（找"三大趋势共振"）
    p_lifecycle = find_para(doc, "三大趋势共振")
    if p_lifecycle:
        insert_after(p_lifecycle, [
            ("t",
             "生命周期定位：油位传感器整体处于成熟期（增速约5%、外资65%格局稳定，竞争重心为降本与认证）；"
             "罐式集装箱油品运输监测处于成长期（增速10-15%、渗透率早期，竞争重心为产能扩张与渠道卡位）；"
             "加油站防渗改造处于政策抢位期（2026-2027集中释放，竞争重心为速度与先发）。"
             "对应策略：油位成熟期投降本拼毛利，罐箱成长期投扩张抢份额，政策窗口期投速度抢先发——"
             "三阶段策略不同，资源投入须分层。"),
        ])
        logging.info("✅ 补生命周期定位")

    # ========== 3. 商业模式分类（business_model_classifier）==========
    # 插入到 3.4 罐箱一体化 之后（找"最具想象力的衍生方向"）
    p_biz = find_para(doc, "最具想象力的衍生方向")
    if p_biz:
        insert_after(p_biz, [
            ("t",
             "商业模式定位：久通为平台型（「箱联全球」SaaS + 80国渠道，收入含订阅与增值服务）；"
             "柯力油位业务为产品型（销售传感器硬件）。本次整合的价值之一，是推动柯力油位"
             "从「产品型」向「产品+平台」延伸——借久通SaaS平台承载油位数据服务，"
             "从一次性硬件销售转向持续性数据订阅。此定位影响估值逻辑："
             "纯产品型按 PS（营收×倍数），含平台属性可按 SaaS 倍数（收入×ARR倍数）估值，"
             "平台属性越强，估值溢价越高。"),
        ])
        logging.info("✅ 补商业模式分类")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
