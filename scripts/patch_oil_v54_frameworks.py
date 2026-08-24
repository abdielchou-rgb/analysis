# -*- coding: utf-8 -*-
"""油位 v5.4：补波特五力 + 中美竞争/宏观传导（圆桌 F-1/F-2 框架缺口）

从 v5.3 复制，补两块框架：
1. 波特五力（插 2.4 潜在进入者后）
2. 中美竞争 + 宏观传导（插 1.5 利润池后）
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging, shutil
logging.basicConfig(level=logging.INFO)

SRC = "油位传感器_行业调研与承接久通生产可行性报告_v5.3.docx"
REPORT = "油位传感器_行业调研与承接久通生产可行性报告_v5.4.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p
        elif kind == "table":
            t = parent.add_table(rows=len(block[1]) + 1, cols=len(block[1][0]),
                                 width=Inches(6.2))
            t.style = "Table Grid"
            for j, h in enumerate(block[1][0]):
                t.cell(0, j).text = h
            for i, row in enumerate(block[1][1:], 1):
                for j, v in enumerate(row):
                    t.cell(i, j).text = str(v)
            prev.addnext(t._tbl)
            prev = t._tbl
            p = parent.add_paragraph("")
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy(SRC, REPORT)
    doc = Document(REPORT)

    # ========== F-1 波特五力（插 2.4 潜在进入者后）==========
    p = find_para(doc, "先发优势可持续期")
    if p:
        insert_after(p, [
            ("h", "2.5 波特五力分析"),
            ("t",
             "油位传感器行业波特五力评估（为中端机会带可切入提供框架依据）："),
            ("table", [
                ["五力", "强度", "评估"],
                ["供应商议价", "高", "波导丝双寡头（爱知制钢/VAC），议价能力强，中游毛利被挤压"],
                ["买方议价", "中", "加油站/危化品客户分散，但认证后转换成本高，议价能力中等"],
                ["替代威胁", "中", "雷达物位计在高端/罐箱场景替代加速，磁致伸缩需差异化（精度/防爆/成本）"],
                ["新进入者", "中", "国产突破者（武汉利又德/博尔森）进入，但认证与渠道构成壁垒"],
                ["现有竞争", "高", "外资五巨头65%锁高端、本土30家小厂锁低端，中端存在空白"],
            ]),
            ("t",
             "五力结论：中端机会带（加油站ATG/罐箱/危化品运输）吸引力来自——供应商议价虽高但通过认证+渠道可对冲、"
             "替代威胁可差异化规避、新进入者受认证壁垒制约、现有竞争在细分存在空白。"
             "五力结构支持『中端可切入』判断，但要求柯力建立认证速度+罐箱渠道两大壁垒。"),
        ])
        logging.info("✅ F-1 波特五力")

    # ========== F-2 中美竞争 + 宏观传导（插 1.5 利润池后）==========
    p = find_para(doc, "1.5 利润池与卡位分析")
    if p:
        insert_after(p, [
            ("h", "1.6 中美竞争与宏观传导"),
            ("t",
             "中美竞争维度：油位传感器核心部件磁致伸缩波导丝由日本爱知制钢、德国VAC主导，"
             "供应链集中度极高。中美科技摩擦背景下，虽波导丝供应方为日德（非直接受美国出口管制对象），"
             "但存在两类风险——①高端传感器材料供应链安全议题升温，各国强化自主可控要求；"
             "②国产替代（武汉利又德/博尔森）受政策鼓励，构成『卡脖子自主可控』的受益方向。"
             "对柯力而言，波导丝断供风险是既有红线（涨价30%触发退出），但国产替代加速是正向期权。"),
            ("t",
             "宏观传导链：油位设备需求与宏观存在传导——原油价格→加油站盈利→防渗改造投资意愿；"
             "利率→基建/危化品投资节奏；国际贸易→罐箱运输量→罐箱液位监测需求。"
             "当前原油价格中枢、贸易周期与政策窗口（2026-2027防渗/SIS改造）共振，"
             "需求端具备支撑；但2028年后政策红利回落需警惕宏观逆风。"),
        ])
        logging.info("✅ F-2 中美竞争+宏观传导")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
