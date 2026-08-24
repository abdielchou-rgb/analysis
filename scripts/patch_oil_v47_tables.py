# -*- coding: utf-8 -*-
"""油位 v4.7：卡点评分转表格 + 类似文字罗列转表格

优化：
1. 波导丝卡点评分（10问20分）→ 表格
2. 组织差距（①技术栈②产线③团队④体系）→ 表格
3. 切换门槛（三个信号gate）→ 保留文字（3项不适合表格，加紧凑）
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.7.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_table_after(anchor_p, headers, rows, width_in=6.2):
    """在锚点段后插入表格。"""
    prev = anchor_p._p
    parent = anchor_p._parent
    t = parent.add_table(rows=len(rows) + 1, cols=len(headers), width=Inches(width_in))
    t.style = "Table Grid"
    # 表头
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
    # 数据
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = str(v)
    prev.addnext(t._tbl)
    # 表后空段
    p = parent.add_paragraph("")
    t._tbl.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.6.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 波导丝卡点评分 → 表格 ==========
    p_cp = find_para(doc, "波导丝卡点评分")
    if p_cp:
        # 替换原文字段为简短结论
        replace_para_text(p_cp,
            "波导丝卡点评分（10问20分制）如下表，合计10/20分，评级「中等偏弱」。"
            "波导丝是行业卡点但非不可突破（国产替代已出现），"
            "柯力进入油位须将「波导丝自研」列为中期战略。")
        # 插入表格
        headers = ["评估维度", "得分", "依据"]
        rows = [
            ["需求不可替代", "2/2", "磁致伸缩无替代方案，波导丝是关键材料"],
            ["供给难扩张", "2/2", "爱知制钢/VAC产能有限，扩产周期长"],
            ["认证设计导入", "2/2", "已有防爆/计量认证体系"],
            ["独家或主供", "1/2", "双寡头，非独家"],
            ["错定价", "0/2", "市场已知，无预期差"],
            ["纯度", "0/2", "非纯卡位标的"],
            ["弹性", "1/2", "需求变化对业绩弹性中等"],
            ["时间窗", "0/2", "窗口期短，先发优势有限"],
            ["护城河", "2/2", "材料专利构成真实壁垒"],
            ["替代风险", "0/2", "国产突破中，替代风险上升"],
        ]
        insert_table_after(p_cp, headers, rows)
        logging.info("✅ 卡点评分转表格")

    # ========== 2. 组织差距 → 表格 ==========
    p_org = find_para(doc, "承接久通油位生产的组织差距")
    if p_org:
        replace_para_text(p_org, "承接久通油位生产的组织差距如下：")
        headers2 = ["维度", "现状", "缺口/投入"]
        rows2 = [
            ["技术栈", "具备传感器制造与认证路径", "磁致伸缩为新增产品线，需认证工程师1-2名（3-6个月）"],
            ["产线", "现有产能可复用", "新增磁致伸缩装配/测试工位，投入200-300万元"],
            ["团队", "无内部候选", "油位产品线硬件负责人（8年经验）为核心缺口，需外部招聘"],
            ["体系", "需建立认证流程", "防爆认证（ATEX/IECEx）12-18个月，为最大时间约束"],
        ]
        insert_table_after(p_org, headers2, rows2)
        logging.info("✅ 组织差距转表格")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
