# -*- coding: utf-8 -*-
"""油位 v4.0：重写为内部整合框架（久通是柯力控股子公司）

核心转变：从"外部合作谈判" → "集团内部整合"。
1. 3.2 合作实质 → 改为"集团内部业务协同"
2. 6.5 交易结构 → 改为"内部整合路径"（删除代工+股权绑定）
3. 新增"控股关系说明"
4. 修正风险框架（少数股东权益替代"换供应商"）
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.0.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v3.2.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 3.2 合作实质 → 内部业务协同 ==========
    p_32 = find_para(doc, "久通提出将油位传感器生产全部交由本公司承接")
    if p_32:
        replace_para_text(p_32,
            "久通物联为柯力传感控股子公司（柯力官网公告「久通成为柯力传感控股子公司」，"
            "纳入集团控股参股体系），本次合作本质是集团内部业务协同，而非外部交易。"
            "久通拥有80余国海关/物流客户网络与罐箱物联网平台，具备油位传感器产品线（JT606X）；"
            "柯力具备制造、认证与规模化能力。将油位传感器生产统筹至柯力制造体系，"
            "是集团内部「制造+渠道」资源整合——既释放久通的渠道与平台价值，"
            "又补齐柯力在液位品类的制造短板。")
        logging.info("✅ 3.2 改内部协同")

    # ========== 2. 删除/改写 双轨"入场券"逻辑（内部无需谈判） ==========
    p_dual = find_para(doc, "两条线的关系可概括为")
    if p_dual:
        replace_para_text(p_dual,
            "两条业务线的关系：代工线（久通现有油位传感器生产，单价约300元/只）是集团内部的存量业务，"
            "柯力承接后可通过制造规模降本、提升良率，将其利润中枢抬升；"
            "中高端线（柯力主导开发、面向加油站/危化品/罐箱的磁致伸缩液位仪，单价约5,000元/台）"
            "是增量主战场，借久通渠道放量。两者同属集团油位业务板块，由柯力统一统筹产能与渠道分配，"
            "不存在「以代工换入场权」的外部谈判逻辑。")
        logging.info("✅ 双轨改内部统筹")

    # ========== 3. 6.5 交易结构 → 内部整合路径 ==========
    p_65 = find_para(doc, "6.5 交易结构选项分析")
    if p_65:
        # 替换整个小节：找到 6.5 标题到 七、结论 之间，改写
        replace_para_text(p_65, "6.5 内部整合路径")
        # 后续段落逐个改写（找到标题后的段落）
        changed = 0
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith("承接久通生产的落地结构有三类可选"):
                replace_para_text(p, "久通已是控股子公司，本方案的本质是集团内部整合，核心路径如下：")
                changed += 1
            elif t.startswith("方案A 纯代工协议"):
                replace_para_text(p, "路径一 制造集中：将油位传感器生产统筹至柯力制造体系，久通聚焦渠道与物联网平台。"
                                     "集团内部产能调配，无需外部签约。")
                changed += 1
            elif t.startswith("方案B 代工+股权绑定"):
                replace_para_text(p, "路径二 渠道统筹：由集团统一规划久通80国渠道在中高端线的推广，"
                                     "明确内部转移定价（成本加成8%-12%，符合独立交易原则），"
                                     "并处理少数股东权益（久通若非全资，需确保交易公允、保护少数股东）。")
                changed += 1
            elif t.startswith("方案C 合资公司"):
                replace_para_text(p, "路径三 平台升级：视业务规模，可考虑将油位业务整合为集团旗下独立事业部或子公司，"
                                     "便于独立核算、考核与未来资本运作。")
                changed += 1
            elif t.startswith("建议：以方案B为主推"):
                replace_para_text(p, "建议：以「制造集中+渠道统筹」为近期路径，"
                                     "业务规模扩大后升级为独立事业部（路径三）。"
                                     "内部整合无「绑定外部合作方」问题，重点在转移定价公允与少数股东权益。")
                changed += 1
        logging.info("✅ 6.5 改内部整合路径（%d段）", changed)

    # ========== 4. 风险框架修正（防换供应商 → 少数股东/内部协同） ==========
    p_risk = find_para(doc, "久通渠道订单未达预期（35%，800万）")
    if p_risk:
        replace_para_text(p_risk,
            "在止损机制基础上，量化主要风险的发生概率与影响金额：政策红利2028回落（概率40%，影响约600万）"
            "最可能发生；渠道订单未达预期（35%，800万）与认证延期（30%，700万）次之；"
            "毛利率低于40%（30%，1,200万）影响最大；波导丝涨价/断供（15%，900万）概率较低但影响大。"
            "内部整合下，「久通更换供应商」风险消除，新增需关注的风险为少数股东权益保护与"
            "内部转移定价公允性（关联交易合规）。")
        logging.info("✅ 风险框架修正")

    # ========== 5. 止损机制修正（删"久通渠道首年订单未达500万收缩海外线"的外部逻辑） ==========
    p_stop = find_para(doc, "久通渠道首年海外订单未达500万元，收缩海外线")
    if p_stop:
        new_stop = p_stop.text.replace("久通渠道首年海外订单未达500万元，收缩海外线",
            "中高端线经久通渠道首年订单未达500万元，收缩海外线投入（保留代工线维持基本盘）")
        replace_para_text(p_stop, new_stop)
        logging.info("✅ 止损机制修正")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
