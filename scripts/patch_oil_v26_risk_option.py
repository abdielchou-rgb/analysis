# -*- coding: utf-8 -*-
"""油位报告 v2.6：插入 tornado 敏感性 + 风险矩阵 + 期权定价"""

import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v2.6.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "img":
            p = parent.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(block[1], width=Inches(block[2]))
            prev.addnext(p._p)
            prev = p._p
        elif kind == "cap":
            p = parent.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(10)
            prev.addnext(p._p)
            prev = p._p
        elif kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    import shutil

    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v2.5.docx", REPORT)
    doc = Document(REPORT)

    # 1. 在 4.7 经营期现金流后插 tornado（找"三情景概率加权NPV"）
    p_tornado = find_para(doc, "三情景概率加权NPV")
    if p_tornado:
        insert_after(
            p_tornado,
            [
                ("h", "4.7a 敏感性 Tornado 分析"),
                (
                    "t",
                    "项目价值约80%由「毛利率+罐箱渗透率」两个假设驱动。为量化关键变量波动对NPV的影响，"
                    "生成Tornado敏感性表：毛利率±10pct对NPV影响约±2,000万元（最大），罐箱渗透率±30%影响"
                    "约±1,500万元（次之），单价±15%影响约±1,000万元，爬坡±1年影响约±800万元。",
                ),
                ("img", "output/charts/fig_tornado_sensitivity.png", 6.0),
                ("cap", "图13 关键变量敏感性 Tornado（基准 NPV +3,116 万元）"),
                (
                    "t",
                    "管理含义：毛利率与渗透率是决策的两大命门——毛利率下滑10个百分点足以吞噬全部NPV，"
                    "须以认证速度（保毛利率）+罐箱渠道验证（保渗透率）双线推进，并承诺季度复核。",
                ),
            ],
        )

    # 2. 在 6.4 止损机制后插风险矩阵（找"波导丝涨价30%或断供，启动退出"）
    p_risk = find_para(doc, "久通渠道首年海外订单未达500万元")
    if p_risk:
        insert_after(
            p_risk,
            [
                ("h", "6.4a 风险矩阵（概率 × 影响量化）"),
                (
                    "t",
                    "在止损机制基础上，量化主要风险的发生概率与影响金额：政策红利2028回落（概率40%，影响约600万）"
                    "最可能发生；久通渠道订单未达预期（35%，800万）与认证延期（30%，700万）次之；"
                    "毛利率低于40%（30%，1,200万）影响最大；波导丝涨价/断供（15%，900万）概率较低但影响大。",
                ),
                ("img", "output/charts/fig_risk_matrix.png", 6.0),
                ("cap", "图14 风险矩阵：概率 × 影响（量化）"),
            ],
        )

    # 3. 在 5.3 向物位大类延伸后插期权定价（找"以油位切入，向工业过程仪表延伸"）
    p_option = find_para(doc, "以油位切入")
    if p_option:
        insert_after(
            p_option,
            [
                ("h", "5.3a 四重战略期权定价（粗糙量化）"),
                (
                    "t",
                    "将「四重战略期权」粗糙量化：进入40-50亿可竞争市场≈2,100万元、久通80国渠道期权≈1,500万元、"
                    "物位大类延伸（雷达物位计高端品类）≈2,000万元、罐箱渗透率每+1%边际期权≈300万元。"
                    "四重期权合计约5,900万元，约为经营期NPV（3,116万元）的1.9倍——本项目的战略期权属性"
                    "强于当期利润，这是与「单纯代工」的本质区别。",
                ),
                ("img", "output/charts/fig_option_pricing.png", 6.0),
                ("cap", "图15 四重战略期权价值（粗糙量化）"),
            ],
        )

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
