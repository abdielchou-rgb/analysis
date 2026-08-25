# -*- coding: utf-8 -*-
"""油位 v2.8：修复下行口径歧义 + 校正正则误报"""

import logging
import shutil

from docx import Document

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v2.8.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def insert_after(anchor_p, block):
    prev = anchor_p._p
    parent = anchor_p._parent
    p = parent.add_paragraph(block)
    prev.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v2.7.docx", REPORT)
    doc = Document(REPORT)

    # 1. 止损机制后补量纲说明（消除 2100 vs 2450 歧义）
    p_stop = find_para(doc, "累计运营投入达2,450万元仍未盈利")
    if p_stop:
        insert_after(
            p_stop,
            "口径说明：2,450万元为「累计运营投入」（含持续经营期投入：设备摊销+团队+认证+运营），"
            "2,100万元为「最坏情形资金敞口」（含设备残值净损后的净现金损失，乐观残值口径）。"
            "两者量纲不同——累计运营投入可能超过初始资金敞口（含持续运营成本），不矛盾；"
            "统一止损基准为「累计运营投入2,450万元触发退出」，最坏敞口2,100万元用于表述股东实际损失上限。",
        )
        print("✅ 已补止损量纲说明")

    # 2. 在结论处补下行边界总结（强化单一决策基准）
    p_concl = find_para(doc, "本项投资具有非对称性")
    if p_concl:
        insert_after(
            p_concl,
            "下行边界统一口径：股东实际损失上限约2,100万元（最坏资金敞口，占2025年归母净利润约6%）；"
            "触发退出基准为累计运营投入2,450万元。两个数字服务于不同决策目的——前者回答「最坏亏多少」，"
            "后者回答「何时止损」，均已在本报告明确。",
        )
        print("✅ 已补结论下行边界")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
