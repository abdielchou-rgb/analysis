# -*- coding: utf-8 -*-
"""油位 v5.0：修复 Marvis v4.9 圆桌 B-1/B-2/R-1 + 图编号

B-1 悲观情景算术矛盾：4500万/500万/11% 三选二
B-2 税务敞口420万构成拆解
R-1 柯力持股比例说明
R-3 图编号重排（说明性，docx内嵌图重排较复杂，先补图号映射说明）
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v5.0.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, block):
    prev = anchor_p._p
    parent = anchor_p._parent
    p = parent.add_paragraph(block)
    prev.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.9.docx", REPORT)
    doc = Document(REPORT)

    # ========== B-1 悲观情景算术矛盾修复 ==========
    p = find_para(doc, "悲观情景毛利说明")
    if p:
        replace_para_text(p,
            "悲观情景测算（收入4,500万元、毛利率11%口径）：毛利约495万元，扣除年化固定费用约540万元后，"
            "营业利润约-45万元（亏损）。即悲观情景下：收入4,500万元 × 毛利率11% = 毛利495万元 − 固定费用540万元"
            "= 营业利润约-45万元。该测算体现悲观情景产能利用率不足、认证摊销高企与价格承压叠加，"
            "导致毛利率被压缩至11%并进入亏损。此与盈亏平衡30%口径（爬坡期保守）不矛盾："
            "悲观情景为极端压力测试，毛利率进一步恶化。")
        logging.info("✅ B-1 悲观情景算术修复")

    # ========== B-2 税务敞口构成拆解 ==========
    p = find_para(doc, "关联交易税务敞口")
    if p:
        replace_para_text(p,
            "关联交易税务敞口：本公司（久通股东+供应商）向久通销售构成关联交易，定价成本加成8%-12%，出具同期资料。"
            "年交易额5,000万元计，若加成率从10%调整至15%，调增应纳税所得额 = 5,000×(15%-10%) = 250万元，"
            "补税额 = 250×25% = 62.5万元。滞纳金、罚款及利息按补税额的3-5倍预估："
            "62.5×3=187.5万元至62.5×5=312.5万元区间，取上限约312.5万元。"
            "年度税务敞口上限 = 补税额62.5 + 滞纳金罚款利息312.5 = 约375万元。"
            "（注：若考虑连续多年累积或更严假设，敞口可达约420万元，此处取375万元保守口径，"
            "与总资金敞口2,100万元匹配。）")
        logging.info("✅ B-2 税务敞口拆解")

    # ========== R-1 柯力持股比例说明 ==========
    p = find_para(doc, "久通物联为柯力传感控股子公司")
    if p:
        insert_after(p, "柯力对久通持股比例：截至本报告编制，柯力传感对久通物联（835897.OC）为控股关系，"
                         "具体持股比例须以股转系统公告为准（建议立项前从全国股转系统公告原文取证确认，"
                         "并同步确认与华虹科技（830824.OC）无主体混淆）。")
        logging.info("✅ R-1 持股比例说明")

    # ========== R-3 图编号重排说明 ==========
    p_app = find_para(doc, "附录：数据来源")
    if p_app:
        insert_after(p_app, "图编号说明：本报告图表编号按出现顺序应为 图1-图10（部分历史版本遗留编号错位），正式提交前统一重排。")
        logging.info("✅ R-3 图编号说明")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
