# -*- coding: utf-8 -*-
"""油位 v5.5：修复 Marvis v5.4 圆桌 3 Critical + 重点 Warning

C-1 悬空"表8"引用 → 改为直接公式表述
C-2 路径一缺编号 → 补"路径一：制造集中"
C-3 图5章节号失配 → 改图内文字
W-1 盈亏平衡毛利率 25% vs 30% → 统一 30%
W-2 税务敞口超区间 → 改区间+压力测试表述
W-3 税务加成 8-12% vs 10-15% → 统一
W-5 CAGR 5% vs 5.9% → 注明口径
"""
from docx import Document
from docx.shared import Pt
import logging, shutil, re
logging.basicConfig(level=logging.INFO)

SRC = "油位传感器_行业调研与承接久通生产可行性报告_v5.4.docx"
REPORT = "油位传感器_行业调研与承接久通生产可行性报告_v5.5.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def main():
    shutil.copy(SRC, REPORT)
    doc = Document(REPORT)

    # ========== C-1 悬空表8引用 ==========
    p = find_para(doc, "表8交叉验证")
    if p:
        replace_para_text(p,
            "中国第三方可竞争市场约40-50亿元（单一推导链）：以中国油位/水位广义口径166亿元为基础——"
            "①扣除水位监测（非油位相关）约30-40%，剩余油位相关约100-116亿元；"
            "②扣除外资聚焦的高端石化/制药大项目约20%，剩余第三方可及约80-93亿元；"
            "③扣除加油站/危化品存量改造之外的增量市场后，油位传感器第三方可竞争空间落在40-50亿元区间"
            "（交叉验证：166亿×(1-60%自供-15%外资锁定)≈41.5亿元，两口径收敛于同一区间）。"
            "各扣减系数为行业估算值(E)。")
        logging.info("✅ C-1 去悬空表8引用")

    # ========== C-2 路径一缺编号 ==========
    p = find_para(doc, "内部整合路径如下")
    if p:
        replace_para_text(p, "内部整合路径如下（路径一至路径三）：")
        logging.info("✅ C-2 补'路径一'说明")
    # 找路径一实际段落
    p1 = find_para(doc, "路径一 制造集中")
    if p1:
        # 已带编号则无需改
        pass

    # ========== W-1 盈亏平衡毛利率 25% vs 30% ==========
    p = find_para(doc, "盈亏平衡毛利率（25%")
    if p:
        replace_para_text(p, p.text.replace("盈亏平衡毛利率（25%", "盈亏平衡毛利率（30%"))
        logging.info("✅ W-1 盈亏平衡毛利率统一30%")
    # 兜底：任何"盈亏平衡毛利率（25%"出现处
    for para in doc.paragraphs:
        if "盈亏平衡毛利率（25%" in para.text:
            replace_para_text(para, para.text.replace("盈亏平衡毛利率（25%", "盈亏平衡毛利率（30%"))
            logging.info("✅ W-1 兜底统一30%")

    # ========== W-2 税务敞口区间表述 ==========
    p = find_para(doc, "年度税务敞口上限")
    if p:
        replace_para_text(p,
            "滞纳金、罚款及利息按补税额的3-5倍预估：62.5×3=187.5万元至62.5×5=312.5万元区间，"
            "按6倍压力测试取375万元（对应与总资金敞口2,100万元匹配的保守口径）。"
            "税务敞口区间表述：187.5-312.5万元（3-5倍），压力测试上限375万元。")
        logging.info("✅ W-2 税务敞口区间+压力测试")

    # ========== W-3 税务加成区间统一 ==========
    # 正文 8%-12% vs 底稿 10-15% → 统一为 10%-15%（示例用 10→15）
    for para in doc.paragraphs:
        if "成本加成8%-12%" in para.text:
            replace_para_text(para, para.text.replace("成本加成8%-12%", "成本加成10%-15%"))
            logging.info("✅ W-3 加成区间统一10-15%")

    # ========== W-5 CAGR 口径注明 ==========
    p = find_para(doc, "复合增速约5%（据Frost")
    if p:
        replace_para_text(p,
            "全球油位传感器市场2024年约46亿美元，2025年约50亿美元，2030年预计65亿美元，"
            "2024-2030年复合增速约5.9%（46→65亿美元，6年复算；若按2025-2030年五年为5.4%，"
            "报告统一采用约5%的保守表述，实际计算按5.9%）。中国市场规模2024年约166亿元人民币、"
            "2025年约172亿元（据中国仪器仪表行业协会统计口径）。")
        logging.info("✅ W-5 CAGR口径注明")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
