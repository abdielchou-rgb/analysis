# -*- coding: utf-8 -*-
"""修改油位 v2.3 报告：补竞争对手对标表 + 替代威胁 + 波导丝修正"""
from docx import Document
from docx.shared import Pt
import logging
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v2.3.docx"


def make_table(doc, rows):
    from docx.shared import Inches
    t = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.5))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.cell(i, j).text = cell
    return t


def insert_after(anchor_p, blocks):
    """在 anchor 段落之后依次插入块（h=小标题/t=正文/table=表格）。"""
    prev = anchor_p._p
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = anchor_p._parent.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            p.alignment = 1
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = anchor_p._parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p
        elif kind == "table":
            tbl = make_table(anchor_p._parent, block[1])
            prev.addnext(tbl._tbl)
            prev = tbl._tbl
            p = anchor_p._parent.add_paragraph("")
            prev.addnext(p._p)
            prev = p._p


def main():
    doc = Document(REPORT)

    # 定位锚点段落
    p22 = p23 = p24 = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if "2.2" in t and "主要竞争" in t:
            p22 = p
        if "2.3" in t and "制胜" in t:
            p23 = p
        if "2.4" in t and "潜在" in t:
            p24 = p
    assert p22 and p23 and p24, f"锚点缺失 p22={p22} p23={p23} p24={p24}"

    # 2.2a 竞争对手量化对标表
    blocks_22 = [
        ("h", "2.2a 竞争对手量化对标（补采 2026-08，据中国工控网2025物/液位仪表市场报告）"),
        ("t",
         "全球高端极（外资）：E+H、VEGA、Emerson、Siemens、Yokogawa、Krohne、Magnetrol、"
         "Honeywell、ABB。共同特征：聚焦石化/制药大项目（投入产出比高），中端合规刚需细分"
         "（加油站ATG/罐箱监测/危化品运输）投入不足。"),
        ("table", [
            ["公司", "起源", "定位", "核心壁垒", "中端市场投入"],
            ["E+H", "瑞士", "高端过程仪表", "SIL/防爆认证全、过程全覆盖", "不足，聚焦大项目"],
            ["VEGA", "德国", "雷达/物位高端", "雷达物位全球领先", "大项目导向"],
            ["Emerson", "美国", "过程仪表平台", "Rosemount全平台", "聚焦大流程"],
            ["Siemens", "德国", "工业自动化", "SITRANS物位线", "项目型"],
            ["Yokogawa", "日本", "过程控制", "石化客户深", "聚焦石化"],
            ["Krohne", "德国", "流量+物位", "物位线完整", "部分细分布局"],
        ]),
        ("t",
         "本土玩家：凡宜科技、青岛澳邦、古大仪表（工控网榜单国产头部）。"),
        ("t",
         "关键洞察：中端机会带（加油站ATG、罐箱监测、危化品运输）确实存在外资空白——"
         "外资因单细分规模有限（投入产出比不足）不聚焦，本土小厂缺防爆/计量认证。"
         "此结构印证了本报告「中端可切入」的核心判断，但各细分市场规模仍需进一步测算。"),
    ]
    insert_after(p22, blocks_22)

    # 重新定位（插入后锚点失效，重新查找）
    p24 = None
    for p in doc.paragraphs:
        if "2.4" in p.text.strip() and "潜在" in p.text.strip():
            p24 = p
            break

    # 2.4a 替代威胁 + 2.4b 波导丝修正
    blocks_24 = [
        ("h", "2.4a 替代威胁（补采）"),
        ("t",
         "雷达物位计是磁致伸缩在罐箱/危化品场景的最大替代威胁——高端市场已逐步转向雷达。"
         "磁致伸缩在罐箱监测的定位须与雷达差异化：高精度（±1mm级）、防爆认证（ATEX/IECEx）、"
         "成本优势（中端价格带）。若不能建立差异化，未来3-5年罐箱细分存在被雷达物位计侵蚀的风险。"),
        ("h", "2.4b 波导丝国产化修正（补采）"),
        ("t",
         "原判断「磁致伸缩波导丝被日企TDK垄断」需修正：调研发现武汉利又德（2024-2025攻克全国产"
         "磁致伸缩液位传感器）、博尔森（纯国产化方案）已实现波导丝国产突破。上游由日本爱知制钢"
         "（Aichi Steel）、德国VAC主导，但国产替代正在发生——自研可行性与时间表可上调。仍保留"
         "波导丝涨价30%/断供触发退出的风险红线（国产良率/一致性需验证），但威胁度下调。"),
    ]
    insert_after(p24, blocks_24)

    doc.save(REPORT)
    print(f"✅ 报告已修改保存: {REPORT}")


if __name__ == "__main__":
    main()
