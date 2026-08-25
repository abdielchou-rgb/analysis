# -*- coding: utf-8 -*-
"""油位报告 v2.7：久通尽调证据 + 组织承接分析 + 数据分级标注"""

import logging
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v2.7.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "img":
            p = parent.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(block[1], width=Inches(block[2]))
            prev.addnext(p._p)
            prev = p._p
        elif kind == "cap":
            p = parent.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(10)
            prev.addnext(p._p)
            prev = p._p
        elif kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v2.6.docx", REPORT)
    doc = Document(REPORT)

    # 1. 3.1 久通概况后补尽调证据（找"年销量约1,000只"）
    p_dd = find_para(doc, "年销量约1,000只")
    if p_dd:
        insert_after(
            p_dd,
            [
                ("h", "3.1a 久通物联尽调证据（工商/公开信息验证，2026-08 补采）"),
                (
                    "t",
                    "经工商信息与公开资料核验：久通物联为新三板挂牌公司（股票代码835897.OC），"
                    "全称「深圳市久通物联科技股份有限公司」，成立于2006年6月13日，注册资本1,436.5万元，"
                    "法定代表人兼董事长陈遵炎（持股约26.5%，380万元）。主营业务为工业化移动与静止燃油消耗"
                    "远程实时监控管理系统、智能车载终端、高精度可截断电容式液位传感器的研发与生产。",
                ),
                (
                    "t",
                    "渠道真实性评估：久通具备油罐车电子锁、油位传感器（JT606X）、集装箱GPS定位等产品线，"
                    "与「80余国海关/物流客户网络」的定位相符（其产品面向海关物流监管场景）。"
                    "但「年销1,000只油位传感器、单价约300元」的建议来源为合作方自述，建议签署合作协议前"
                    "以久通公开财报/对账单验证真实销量与渠道客户名单。",
                ),
            ],
        )

    # 2. 3.3 能力互补后补组织承接分析（找"缺海外油品运输场景入口"）
    p_org = find_para(doc, "缺海外油品运输场景入口")
    if p_org:
        insert_after(
            p_org,
            [
                ("h", "3.3a 组织承接能力分析（差距量化）"),
                (
                    "t",
                    "承接久通油位生产需评估组织差距：①技术栈——本公司具备传感器制造与认证路径，但"
                    "磁致伸缩液位仪为新增产品线，需补充防爆/计量认证工程师1-2名（招聘周期3-6个月）；"
                    "②产线——现有产能可复用，但需新增磁致伸缩装配/测试工位（投入约200-300万元）；"
                    "③团队——前置条件「油位产品线硬件负责人8年以上经验」为核心缺口，目前无内部候选，"
                    "需外部招聘或与久通技术团队共建；④体系——需建立防爆认证（ATEX/IECEx）流程，"
                    "认证周期12-18个月为最大时间约束。",
                ),
            ],
        )

    # 3. 附录前补数据分级标注说明（找"附录：数据来源"）
    p_app = find_para(doc, "附录：数据来源")
    if p_app:
        insert_after(
            p_app,
            [
                ("h", "附录A：数据分级标注"),
                (
                    "t",
                    "本报告数据按可信度分级（R87 体系）：(A)=实际值（公司披露/年报/公告）；"
                    "(E)=估算值（基于模型/假设/行业调研推算，如波导丝国产化率25-30%、"
                    "加油站执行率5-8%、单站价值5-10万元）；(F)=预测值（前瞻判断，如2030年市场规模65亿美元）；"
                    "(B)=基准值（同业/行业对标）。E级数据为决策参考，签署协议前需以客户访谈/供应商尽调升级为A级。",
                ),
            ],
        )

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
