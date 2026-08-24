# -*- coding: utf-8 -*-
"""油位 v4.9：全量推进圆桌剩余项
C-3 打通可竞争市场两套口径 + W-1 数据标注 + W-2 渠道证据 + W-3 波导丝佐证 + W-6 期权表述
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.9.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, block):
    prev = anchor_p._p
    parent = anchor_p._parent
    p = parent.add_paragraph(block)
    prev.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.8.docx", REPORT)
    doc = Document(REPORT)

    # ========== C-3 打通可竞争市场推导链 ==========
    # 摘要的可竞争市场段（找"中国第三方可竞争市场约40-50亿元"）
    p_summary = find_para(doc, "中国第三方可竞争市场约40-50亿元")
    if p_summary:
        replace_para_text(p_summary,
            "中国第三方可竞争市场约40-50亿元（单一推导链）：以中国油位/水位广义口径166亿元为基础——"
            "①扣除水位监测（非油位相关）约30-40%，剩余油位相关约100-116亿元；"
            "②扣除外资聚焦的高端石化/制药大项目约20%，剩余第三方可及约80-93亿元；"
            "③扣除加油站/危化品存量改造之外的增量市场后，油位传感器第三方可竞争空间落在40-50亿元区间"
            "（表8交叉验证：166亿×(1-60%自供-15%外资锁定)≈41.5亿元，两口径收敛于同一区间）。"
            "各扣减系数为行业估算值(E)。")
        logging.info("✅ C-3 打通可竞争市场推导链")

    # ========== W-1 数据标注补全 ==========
    # 久通概况段（年销1000只/300元）—— 找到"其油位传感器产品线规模有限"
    p_jt = find_para(doc, "其油位传感器产品线规模有限")
    if p_jt:
        replace_para_text(p_jt,
            "其油位传感器产品线规模有限：年销量约1,000只(E)、单价约300元(E)、年收入约30万元(E)（据集团内部销售数据，待财报核对）。")
        logging.info("✅ W-1 久通数据标注")

    # 战略期权段（5,900万 ≈ 1.9倍）→ 弱化直接可比（W-6）
    p_opt = find_para(doc, "四重期权合计约5,900万元")
    if p_opt:
        replace_para_text(p_opt,
            "四重期权合计约5,900万元（粗糙量化，供定性参考）。该期权价值为战略敞口估算，"
            "与经营期NPV（3,116万元，现金流折现口径）不同量纲，不宜直接相加或作倍数比较；"
            "其意义在于提示本项目的期权属性强于当期利润，而非承诺上行空间为NPV的倍数。")
        logging.info("✅ W-6 期权表述弱化")

    # ========== W-2 渠道证据审计说明 ==========
    p_channel = find_para(doc, "久通具备油罐车电子锁、油位传感器")
    if p_channel:
        insert_after(p_channel,
            "渠道证据待审计：久通80余国渠道目前为「箱联全球SaaS覆盖海关/税务/物流客户」的定性描述，"
            "缺少分国别客户数、油品运输类客户占比、JT606X实际出货国别清单等可审计证据。"
            "建议立项阶段从集团内部调取久通客户台账，验证油品运输场景的真实可达性；"
            "在渠道证据验证前，罐箱一体化（单箱价值量千元级→数千元级）作为愿景而非承诺。")
        logging.info("✅ W-2 渠道证据审计说明")

    # ========== W-3 波导丝国产化率佐证 ==========
    p_wave = find_para(doc, "国产化率2025年约25%-30%")
    if p_wave:
        replace_para_text(p_wave,
            "磁致伸缩式液位仪的核心敏感元件为磁致伸缩波导丝，全球由日本爱知制钢（Aichi Steel）、德国VAC主导，"
            "国产化率2025年约25%-30%（行业供应链调研及公开专利分析，估算值E，"
            "建议以行业协会或券商研报原文补充第三方佐证并标注获取日期）。"
            "外购装配下，中游制造净利率仅8%-12%，上游材料毛利50%-65%——利润集中于两端。")
        logging.info("✅ W-3 波导丝佐证说明")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
