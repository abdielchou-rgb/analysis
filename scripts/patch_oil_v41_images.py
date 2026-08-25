# -*- coding: utf-8 -*-
"""油位 v4.1：修复空引用/截断段落 + 嵌入缺失图表

修复：
1. [24] "上述「踩踏」风险" 截断 → 恢复完整政策回落情景段
2. [25] "加油站执行率" 截断 → 恢复完整数据来源段
3. 图1 竞争格局矩阵 → 嵌入 fig_bcg_competitive.png
4. 图5 收入预测三情景 → 嵌入 fig_fin_revenue_scenarios.png
5. 图13/14/15 → 检查并重新嵌入
"""

import logging
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.1.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_img_after(anchor_p, img_path, caption, width=6.0):
    """在锚点段后插入图片 + 图注。"""
    prev = anchor_p._p
    parent = anchor_p._parent
    # 图注（放在图前？报告惯例是图注在图上或图下）
    # 图
    p_img = parent.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(img_path, width=Inches(width))
    prev.addnext(p_img._p)
    prev = p_img._p
    # 图注
    p_cap = parent.add_paragraph()
    run = p_cap.add_run(caption)
    run.bold = True
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.size = Pt(10)
    prev.addnext(p_cap._p)
    prev = p_cap._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.0.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 恢复截断段落 ==========
    p_ca = find_para(doc, "上述「踩踏」风险")
    if p_ca and len(p_ca.text.strip()) < 30:
        replace_para_text(
            p_ca,
            "上述「踩踏」风险应纳入财务测算：悲观情景假设2028年后政策红利回落，"
            "加油站防渗替换与SIS改造需求增速降至零甚至小幅收缩，对应悲观NPV约-1,100万元。"
            "产能规划保持弹性、不做重资产投入即是为该情景预留缓冲。",
        )
        logging.info("✅ 恢复「踩踏」风险段")

    p_rate = find_para(doc, "加油站执行率")
    if p_rate and len(p_rate.text.strip()) < 30:
        replace_para_text(
            p_rate,
            "加油站执行率（全国尚有5%-8%未完成）、单站价值5-10万元、SIS改造43-52亿元均为行业测算口径估算值(E)，"
            "待客户访谈与供应商尽调验证后升级为实际值。波导丝国产化率25%-30%同样为估算值(E)。",
        )
        logging.info("✅ 恢复「加油站执行率」段")

    # ========== 2. 嵌入图1 竞争格局矩阵 ==========
    p_fig1 = find_para(doc, "图1 油位传感器竞争格局矩阵")
    if p_fig1:
        # 检查后段是否有图
        has_img = len(p_fig1._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")) > 0
        if not has_img:
            insert_img_after(p_fig1, "output/charts/fig_bcg_competitive.png", "图1 油位传感器竞争格局矩阵")
            logging.info("✅ 嵌入图1")

    # ========== 3. 嵌入图5 收入三情景 ==========
    p_fig5 = find_para(doc, "图5 收入预测三情景")
    if p_fig5:
        has_img = len(p_fig5._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")) > 0
        if not has_img:
            insert_img_after(p_fig5, "output/charts/fig_fin_revenue_scenarios.png", "图5 收入预测三情景")
            logging.info("✅ 嵌入图5")

    # ========== 4. 检查并嵌入图13/14/15 ==========
    fig_map = [
        ("图13 关键变量敏感性", "output/charts/fig_tornado_sensitivity.png"),
        ("图14 风险矩阵", "output/charts/fig_risk_matrix.png"),
        ("图15 四重战略期权价值", "output/charts/fig_option_pricing.png"),
    ]
    for frag, path in fig_map:
        p = find_para(doc, frag)
        if p:
            has_img = len(p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")) > 0
            if not has_img:
                insert_img_after(p, path, p.text.strip(), width=6.0)
                logging.info("✅ 嵌入 %s", frag[:20])

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
