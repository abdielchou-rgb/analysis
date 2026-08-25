# -*- coding: utf-8 -*-
"""油位 v5.1：图编号重排 + 测算结果标注 + 底稿清单

R-3 图编号重排：图0/0b/1/12/13/5/15/14 → 图1-图8 连续
R-2 测算结果标注：NPV/IRR/盈亏平衡等补 (A)/(E)/(F)
R-4 测算底稿清单：附录列出
"""

import logging
import shutil

from docx import Document
from docx.shared import Inches

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v5.1.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v5.0.docx", REPORT)
    doc = Document(REPORT)

    # ========== R-3 图编号重排 ==========
    # 按出现顺序映射：图0→图1, 图0b→图2, 图1→图3, 图12→图4, 图13→图5, 图5→图6, 图15→图7, 图14→图8
    fig_map = {
        "图0 油位传感器市场与承接机会总览": "图1 油位传感器市场与承接机会总览",
        "图0b 承接久通机会逻辑链：以代工为起点换取「制造+渠道」协同战略入口": "图2 承接久通机会逻辑链：以代工为起点换取「制造+渠道」协同战略入口",
        "图1 油位传感器竞争格局矩阵": "图3 油位传感器竞争格局矩阵",
        "图12 推导总桥：从市场总量到NPV": "图4 推导总桥：从市场总量到NPV",
        "图13 关键变量敏感性 Tornado（基准 NPV +3,116 万元）": "图5 关键变量敏感性 Tornado（基准 NPV +3,116 万元）",
        "图5 收入预测三情景": "图6 收入预测三情景",
        "图15 四重战略期权价值（粗糙量化）": "图7 四重战略期权价值（粗糙量化）",
        "图14 风险矩阵：概率 × 影响（量化）": "图8 风险矩阵：概率 × 影响（量化）",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in fig_map:
            replace_para_text(p, fig_map[t])
    logging.info("✅ R-3 图编号重排（图1-图8）")

    # 删除图编号说明段（已重排，不再需要"正式提交前统一重排"）
    p_note = find_para(doc, "图编号说明：本报告图表编号按出现顺序应为")
    if p_note:
        p_note._element.getparent().remove(p_note._element)
        logging.info("✅ 删除图编号说明段")

    # ========== R-2 测算结果标注 ==========
    # NPV/IRR 段
    p_npv = find_para(doc, "经营期NPV = +3,116万元为唯一决策主指标")
    if p_npv:
        replace_para_text(
            p_npv,
            "经营期NPV = +3,116万元(F，测算值)为唯一决策主指标。含终值NPV（永续增长法约3.2亿元、"
            "8倍EBITDA法约2.2亿元）仅作敏感性参考，不作为主指标——"
            "终值依赖永续增长率/退出倍数等强假设，对1,700万元级项目引入过高杠杆，"
            "决策统一以经营期NPV为准以保持口径唯一。",
        )
        logging.info("✅ R-2 NPV标注")
    # IRR 段
    p_irr = find_para(doc, "经营期IRR约57%（该高IRR主要因投资基数小")
    if p_irr:
        replace_para_text(
            p_irr,
            "经营期IRR约57%(F，测算值)——该高IRR主要因投资基数小、期限短，不代表回报确定性高，"
            "决策应关注绝对NPV而非IRR。",
        )
        logging.info("✅ R-2 IRR标注")
    # 盈亏平衡段
    p_be = find_para(doc, "盈亏平衡月收入 = 45 ÷ (1-0.70)")
    if p_be:
        replace_para_text(
            p_be,
            "盈亏平衡月收入 = 45 ÷ (1-0.70) = 150万元/月(F，测算值)，对应年化约1,800万元，预计2027年第二季度达成。",
        )
        logging.info("✅ R-2 盈亏平衡标注")

    # ========== R-4 测算底稿清单 ==========
    p_app = find_para(doc, "附录：数据来源")
    if p_app:
        # 在附录后插入底稿清单表
        prev = p_app._p
        parent = p_app._parent
        p_h = parent.add_paragraph("测算底稿清单")
        p_h.runs and None
        run = p_h.add_run("测算底稿清单")
        run.bold = True
        prev.addnext(p_h._p)
        prev = p_h._p
        tbl = parent.add_table(rows=6, cols=3, width=Inches(6.2))
        tbl.style = "Table Grid"
        data = [
            ["核心数字", "底稿", "假设/来源"],
            ["经营期NPV 3,116万", "contract_manufacturing 模型", "WACC 9.5%, 5年经营期"],
            ["盈亏平衡月收入150万", "固定费用+毛利率模型", "固定月费用45万, 毛利率30%"],
            ["三情景收入/NPV", "三情景测算表", "乐观30%/基准50%/悲观20%"],
            ["税务敞口375万", "关联交易定价测算", "成本加成10-15%, 补税额62.5万×3-5倍"],
            ["最坏敞口2,100万", "运营损失+税务敞口合计", "运营1,700万+税务375-420万"],
        ]
        for i, row in enumerate(data):
            for j, v in enumerate(row):
                tbl.cell(i, j).text = v
        prev.addnext(tbl._tbl)
        logging.info("✅ R-4 测算底稿清单")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
