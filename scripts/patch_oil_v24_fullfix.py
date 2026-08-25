# -*- coding: utf-8 -*-
"""修改油位 v2.4 报告：
1. 图0 拆成 图0a(市场总览) + 图0b(机会逻辑链)，全幅大字
2. 图12 推导总桥重画 + 新增文字论述段
3. 市场调研4处修订（可竞争市场推导/政策回落/增速口径/来源加固）
"""

import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("patch")

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v2.5.docx"


def find_para(doc, text_frag):
    for p in doc.paragraphs:
        if text_frag in p.text:
            return p
    return None


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "img":
            p = parent.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(block[1], width=Inches(block[2]))
            prev.addnext(p._p)
            prev = p._p
        elif kind == "cap":
            p = parent.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(10)
            prev.addnext(p._p)
            prev = p._p
        elif kind == "h":
            p = parent.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def replace_image_in_para(para, img_path, width_in=6.5):
    """替换段落内嵌图片。"""
    from docx.oxml.ns import qn

    drawings = para._p.findall(".//" + qn("w:drawing"))
    for d in drawings:
        para._p.remove(d)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document(REPORT)

    # ========== 1. 图0 拆两张 ==========
    p_fig0 = find_para(doc, "图0 油位传感器市场与承接机会总览")
    if p_fig0:
        # 找到图0所在段落（标题前或后？标题段落本身可能有图）
        # 替换标题段落为两图+标题
        # 在标题段落后插入两图
        insert_after(
            p_fig0,
            [
                ("cap", "图0a 油位传感器市场规模总览（全球46→65亿美元 CAGR≈5%；中国166→172亿元）"),
                ("img", "output/charts/fig0a_market_overview.png", 6.5),
                ("cap", "图0b 承接久通机会逻辑链：以代工为起点换取「制造+渠道」协同战略入口"),
                ("img", "output/charts/fig0b_opportunity_chain.png", 6.5),
            ],
        )
        # 删除原标题段落（避免重复）——保留标题但改文字
        p_fig0.clear()
        run = p_fig0.add_run("图0 油位传感器市场与承接机会总览")
        run.bold = True
        p_fig0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========== 2. 图12 推导总桥重画 + 论述 ==========
    p_fig12 = find_para(doc, "图12 推导总桥")
    if p_fig12:
        # 替换图12标题段落里的图（若有），并插入新图 + 论述
        insert_after(
            p_fig12,
            [
                ("img", "output/charts/fig12_bridge.png", 6.5),
                ("h", "推导总桥论述（图12 展开）"),
                (
                    "t",
                    "第1步 市场规模：全球油位传感器市场2024年46亿美元→2030年65亿美元（CAGR≈5%）。"
                    "中国油位/水位广义口径166→172亿元。扣除外资聚焦的高端大项目与水位监测，"
                    "中国第三方可竞争市场约40-50亿元。",
                ),
                (
                    "t",
                    "第2步 目标细分：选择罐箱油品运输监测（全球6-9亿美元，占油位市场约15%，增速10-15%）"
                    "作为切入点——它是增速最高的油位应用场景，且外资投入不足、本土小厂缺认证，存在机会空白。",
                ),
                (
                    "t",
                    "第3步 渠道可及：依托久通80余国海关/物流客户网络，目标渗透率驱动可及收入。"
                    "罐箱「锁+液位」一体化使单箱价值量从千元级提升至数千元级。",
                ),
                (
                    "t",
                    "第4步 收入测算（自下而上）：海外线（经久通渠道）+国内线（加油站/危化品），"
                    "三情景收入曲线见正文。基准情景下2027年Q2实现盈亏平衡（月收入150万元）。",
                ),
                (
                    "t",
                    "第5步 毛利率：海外综合毛利率40%-50%（公司2025年海外毛利率56.58%为锚，新品类导入期取40-50%）。"
                    "毛利率40-50%下贡献毛利率约30%，对应盈亏平衡月收入150万元（固定月费用45万元÷(1-0.70)）。",
                ),
                (
                    "t",
                    "第6步 净贡献与NPV：自由现金流折现（WACC 9.5%），经营期NPV +3,116万元，IRR约57%。"
                    "三情景概率加权NPV约+3,200万元。",
                ),
                (
                    "t",
                    "第7步 敏感性：项目价值约80%由「毛利率+罐箱渗透率」两个假设驱动——"
                    "须对这两项做季度复核与压力测试（±30%波动）。",
                ),
            ],
        )

    # ========== 3. 市场调研4处修订 ==========
    # 3.1 增速口径：在1.1节补一句说明 2024→2025 与 CAGR 差异
    p_size = find_para(doc, "全球油位传感器市场2024年约46亿美元")
    if p_size:
        insert_after(
            p_size,
            [
                (
                    "t",
                    "口径说明：2024→2025年从46亿增至50亿美元（+8.7%），高于2024-2030年CAGR≈5%，"
                    "原因是2024年受下游去库存影响基数偏低、2025年合规驱动集中释放；"
                    "2030年前增速回落至约5%的稳态水平。",
                ),
            ],
        )
    # 3.2 政策回落纳入悲观情景 + 3.4 来源加固：在1.2节后补
    p_policy = find_para(doc, "2028年后政策红利或回落")
    if p_policy:
        insert_after(
            p_policy,
            [
                (
                    "t",
                    "情景修正：上述「踩踏」风险应纳入财务测算——悲观情景假设2028年后政策红利回落，"
                    "加油站防渗替换与SIS改造需求增速降至零甚至小幅收缩，对应悲观NPV约-1,100万元。"
                    "产能规划保持弹性、不做重资产投入即是为该情景预留缓冲。",
                ),
                (
                    "t",
                    "数据来源说明：加油站执行率（全国尚有5%-8%未完成）、单站价值5-10万元、SIS改造43-52亿元"
                    "均来自战略部测算口径，为估算值(E)；波导丝国产化率25%-30%来自行业供应链调研及公开专利分析，"
                    "亦为估算值(E)，待客户访谈/供应商尽调验证后升级为verified。",
                ),
            ],
        )
    # 3.3 可竞争市场推导：在1.1后补推导过程
    p_note = find_para(doc, "不可直接相除")
    if p_note:
        insert_after(
            p_note,
            [
                (
                    "t",
                    "中国可竞争市场推导（40-50亿元）：从中国油位/水位广义口径166亿元出发——"
                    "①扣除水位监测（非油位相关）约30-40%；②扣除外资聚焦的高端石化/制药大项目约20%；"
                    "③扣除加油站/危化品存量改造以外的增量市场后，油位传感器第三方可竞争市场落在40-50亿元区间。"
                    "该推导为战略部测算口径，需以细分行业报告交叉验证。",
                ),
            ],
        )

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
