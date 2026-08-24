# -*- coding: utf-8 -*-
"""油位 v3.0：修复 300元代工 vs 5000元中高端 双轨逻辑 + 三情景指代

问题：
1. 久通给的是300元油位传感器，但测算用5000元中高端磁致伸缩——没讲清两条业务线如何承接
2. "三情景收入曲线见正文"——推导总桥里正文已推导，指代冗余
"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v3.0.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, block):
    prev = anchor_p._p
    parent = anchor_p._parent
    p = parent.add_paragraph(block)
    prev.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v2.9.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 3.2 合作实质后补"双轨业务定位" ==========
    p_32 = find_para(doc, "对本公司，真正的机会不是这30万元订单")
    if p_32:
        insert_after(p_32,
            "需要明确的是，本方案涉及两条定位不同的业务线，其关系如下："
            "其一，「代工线」——承接久通现有油位传感器生产（单价约300元/只，年销量约1,000只，"
            "对应年收入约30万元）。该线利润极薄，其价值不在当期收入，而在三点：①锁定与久通的合作"
            "关系（避免被其他代工厂切入）；②获得久通80余国渠道的进入许可；③在合作中掌握久通客户"
            "的真实需求与订单节奏。代工线本质是「入场券」，非利润来源。"
            "其二，「中高端线」——本公司依托自身制造与认证能力，开发面向加油站防渗改造、"
            "危化品储罐安全仪表、罐箱油品运输监测的磁致伸缩液位仪，单价3,000-8,000元/台"
            "（本方案取4,500-6,500元，中位5,000元/台）。该线才是本方案收入与NPV测算的主体，"
            "其逻辑是：以代工线切入久通体系 → 借久通海外渠道销售中高端产品 → "
            "罐箱「锁+液位」一体化提升单箱价值量。"
            "两条线的关系可概括为「以300元代工换5000元中高端的入场权」——"
            "测算所依据的5,000元/台，对应的是本公司主导开发的中高端磁致伸缩产品，"
            "而非久通现售的300元低端产品。此定位是本方案财务测算成立的前提。")
        logging.info("✅ 补双轨业务定位")

    # ========== 2. 推导总桥第4步 修正指代 ==========
    p_step4 = find_para(doc, "第4步 收入测算（自下而上）")
    if p_step4:
        replace_para_text(p_step4,
            "第4步 收入测算：收入主体为中高端磁致伸缩液位仪（单价约5,000元/台），"
            "分海外线（经久通渠道）与国内线（加油站/危化品）两条路径自下而上测算，"
            "三情景收入曲线见图5。基准情景下2027年Q2实现盈亏平衡（月收入150万元，"
            "对应月销约300台中高端液位仪）。")
        logging.info("✅ 修正推导总桥第4步指代")

    # ========== 3. 4.2 单价假设 补逻辑链 ==========
    p_42 = find_para(doc, "磁致伸缩液位仪价格带3,000-8,000元/台")
    if p_42:
        replace_para_text(p_42,
            "磁致伸缩液位仪价格带3,000-8,000元/台。本方案测算主体为本公司主导开发的中高端磁致伸缩产品，"
            "而非久通现售的300元低端产品（低端产品仅作代工入场券，见3.2业务定位）。"
            "依据竞争定位（中高端磁致伸缩、避开低端红海、不拼最高端防爆），对应价格带4,500-6,500元，"
            "取中位5,000元/台。该单价对应的是通过防爆/计量认证、面向加油站与罐箱监测的合规产品，"
            "与300元低端产品在认证、精度、应用场景上均不同。")
        logging.info("✅ 修正4.2单价逻辑")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
