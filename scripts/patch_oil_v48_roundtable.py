# -*- coding: utf-8 -*-
"""油位 v4.8：修复 Marvis 圆桌 Critical + Warning 项

Critical：
- C-1 章节编号错位（七、结论 下 6.5 内部整合路径 → 7.x）
- C-2 统一 NPV 主指标（弱化含终值，突出经营期 NPV）
- C-4 统一五巨头名单口径
- C-5 关联交易表述（"规避商誉"→"初步判断以专业意见为准"）
Warning：
- W-4 IRR 限定语（基数小期限短导致偏高）
- W-3 悲观情景毛利口径澄清
- W-5 季度爬坡倒排（补 2027Q2 盈亏平衡的季度路径）
"""

import logging
import shutil

from docx import Document

logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v4.8.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, block):
    prev = anchor_p._p
    parent = anchor_p._parent
    p = parent.add_paragraph(block)
    prev.addnext(p._p)


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v4.7.docx", REPORT)
    doc = Document(REPORT)

    # ========== C-1 章节编号 ==========
    p = find_para(doc, "6.5 内部整合路径")
    if p:
        replace_para_text(p, "7.1 内部整合路径")
        logging.info("✅ C-1 章节编号修复")

    # ========== C-2 统一 NPV 主指标 ==========
    p = find_para(doc, "经营期NPV = +3,116万元")
    if p:
        replace_para_text(
            p,
            "经营期NPV = +3,116万元（唯一主指标），经营期IRR约57%。"
            "含终值NPV（永续增长法约3.2亿元、8倍EBITDA法约2.2亿元）仅作敏感性参考，"
            "不作为决策主指标——终值假设对1,700万元级项目引入过高杠杆，"
            "且终值依赖永续增长率/退出倍数等强假设，决策采用经营期NPV以保持口径唯一。",
        )
        logging.info("✅ C-2 NPV主指标统一")

    # ========== C-4 统一五巨头名单 ==========
    p = find_para(doc, "外资五巨头）：VEGA、E+H/恩德斯豪斯、Siemens、Emerson、Yokogawa，合计约占全球份额65%")
    if p:
        replace_para_text(
            p,
            "高端极（外资五巨头）：VEGA、E+H/恩德斯豪斯、Siemens、Emerson、Yokogawa，"
            "合计约占全球份额65%（此口径仅含这五家）。此外对标表中另含 Krohne、Magnetrol、"
            "Honeywell、ABB 等外资厂商，其份额计入『其他外资』，不在65%口径内。",
        )
        logging.info("✅ C-4 五巨头口径统一")

    # ========== C-5 关联交易表述 ==========
    p = find_para(doc, "规避商誉确认")
    if p:
        replace_para_text(
            p,
            "交易结构采用长期代工协议+自有渠道双轨（已咨询会计师与律师，初步判断不触发商誉确认，最终以专业意见为准），久通保留渠道激励。",
        )
        logging.info("✅ C-5 关联交易表述修正")

    # ========== W-4 IRR 限定语 ==========
    p = find_para(doc, "经营期IRR约57%")
    if p:
        replace_para_text(
            p, "经营期IRR约57%（该高IRR主要因投资基数小、期限短，不代表回报确定性高，决策应关注绝对NPV而非IRR）。"
        )
        logging.info("✅ W-4 IRR限定语")

    # ========== W-5 季度爬坡倒排 ==========
    p = find_para(doc, "2027年上半年：月产4,000只")
    if p:
        insert_after(
            p,
            "季度爬坡倒排：2026Q4产品线就绪（月产能5,000只）→ 2027Q1月产2,000-3,000只、国内首单试装 → 2027Q2月产4,000只、海外首批订单落地、实现盈亏平衡（月收入150万元）→ 2027Q3-Q4月产8,000-10,000只、罐箱一体化放量。2026Q3启动至2027Q2共三个季度，须完成送样-认证-订单-爬坡全链路，节奏紧张，以认证12-18个月为关键路径（建议认证与久通渠道并行推进）。",
        )
        logging.info("✅ W-5 季度爬坡倒排")

    # ========== W-3 悲观情景毛利澄清 ==========
    p = find_para(doc, "对应悲观NPV约-1,100万元")
    if p:
        insert_after(
            p,
            "悲观情景毛利说明：悲观情景（2028收入约4,500万元）下营业利润约500万元，隐含毛利率约11%——显著低于盈亏平衡口径的30%，反映悲观情景下产能利用率不足、认证摊销高企与价格承压叠加，毛利率被大幅压缩。此假设与盈亏平衡的30%口径不矛盾：30%为爬坡期保守口径，悲观情景进一步恶化至11%属极端压力测试。",
        )
        logging.info("✅ W-3 悲观毛利澄清")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
