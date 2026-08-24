# -*- coding: utf-8 -*-
"""油位 v3.2：补交易结构三方案 + 统一毛利率口径（圆桌高盛/四大建议）"""
from docx import Document
from docx.shared import Pt
import logging, shutil
logging.basicConfig(level=logging.INFO)

REPORT = "output/油位传感器_行业调研与承接久通生产可行性报告_v3.2.docx"


def find_para(doc, frag):
    for p in doc.paragraphs:
        if frag in p.text:
            return p
    return None


def replace_para_text(para, new_text):
    if para.runs:
        fmt = para.runs[0].font
        para.clear()
        run = para.add_run(new_text)
        run.font.size = fmt.size
        run.bold = fmt.bold
    else:
        para.clear()
        para.add_run(new_text)


def insert_after(anchor_p, blocks):
    prev = anchor_p._p
    parent = anchor_p._parent
    for block in blocks:
        kind = block[0]
        if kind == "h":
            p = parent.add_paragraph()
            p.add_run(block[1]).bold = True
            prev.addnext(p._p)
            prev = p._p
        elif kind == "t":
            p = parent.add_paragraph(block[1])
            prev.addnext(p._p)
            prev = p._p


def main():
    shutil.copy("output/油位传感器_行业调研与承接久通生产可行性报告_v3.1.docx", REPORT)
    doc = Document(REPORT)

    # ========== 1. 交易结构三方案（高盛建议）——插在结论前 ==========
    p_concl = find_para(doc, "七、结论")
    if p_concl:
        insert_after(p_concl, [
            ("h", "6.5 交易结构选项分析"),
            ("t",
             "承接久通生产的落地结构有三类可选，控制权、现金流与风险各不相同："),
            ("t",
             "方案A 纯代工协议：仅签署长期代工协议。优点：法律关系简单、不涉及股权；"
             "缺点：绑定弱，久通可随时更换供应商，渠道价值未锁定。"),
            ("t",
             "方案B 代工+股权绑定：在代工协议基础上，参与久通定增或受让其老股（久通为新三板挂牌公司，"
             "有公开交易与财报）。优点：以股东身份锁定渠道，享有久通成长收益；"
             "缺点：构成关联交易（供应商+股东双重身份），须履行上市公司关联交易合规流程。"),
            ("t",
             "方案C 合资公司：双方合资成立油位产品公司，柯力控股。优点：资产隔离、权责清晰、"
             "可独立融资与退出；缺点：设立与治理成本高，久通渠道资源转移需谈判。"),
            ("t",
             "建议：以方案B为主推（代工+股权绑定，控制权与渠道锁定兼顾），"
             "方案C作为久通渠道验证后的升级路径。方案A仅作过渡。"),
        ])
        logging.info("✅ 补交易结构三方案")

    # ========== 2. 统一毛利率口径（四大建议）——修正4.3 ==========
    p_margin = find_para(doc, "海外毛利率40%-50%")
    if p_margin:
        replace_para_text(p_margin,
            "海外中高端产品毛利率40%-50%（公司2025年海外综合毛利率56.58%为锚，新品类导入期议价较弱取40-50%）。"
            "需要说明：4.6节盈亏平衡计算采用毛利率30%（保守口径，对应可变成本率70%），"
            "两者口径差异源于——4.3的40-50%是满产稳态下的目标毛利率，4.6的30%是爬坡期（产能利用率不足、"
            "固定成本摊销高）的保守毛利率。盈亏平衡点150万元/月（月销约300台）按30%保守口径测算，"
            "达成稳态后随毛利率回升至40-50%，净贡献将显著改善。")
        logging.info("✅ 统一毛利率口径")

    # ========== 3. 双轨资源配置 + 切换gate（MBB建议）——插在3.2后 ==========
    p_dual = find_para(doc, "两条线的关系可概括为")
    if p_dual:
        insert_after(p_dual, [
            ("t",
             "资源配置：代工线投入上限——代工线仅占用现有产线约10%产能与1名工艺工程师，"
             "总投资不超过200万元（含产线调试与送样），其定位是渠道入场券而非利润来源，"
             "投入严格受限。中高端线占用主要资源（认证、研发、产线），是投入与考核的主体。"),
            ("t",
             "切换门槛：从代工线切换到中高端线（经久通渠道放量）以三个信号为gate——"
             "①取得3家以上海外意向客户（含罐箱运营方）；②久通渠道首年订单承诺达500万元；"
             "③中高端产品通过防爆/计量认证。三项满足后方启动中高端线量产，否则维持代工线试产。"),
        ])
        logging.info("✅ 补双轨资源配置+切换gate")

    doc.save(REPORT)
    print("✅ 报告已修改保存:", REPORT)


if __name__ == "__main__":
    main()
