# -*- coding: utf-8 -*-
"""MD → DOCX 专业排版转换器（python-docx）

对标顶级机构报告排版：
- 标题层级：H1=18pt加粗/居中，H2=14pt加粗，H3=12pt加粗
- 表格：表头加底纹+加粗，内容9.5pt，Table Grid边框
- 页边距：2.5cm
- 正文：11pt，1.5倍行距，首行缩进
- 图片：居中，宽14cm，图题加粗
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cn_font(run, name="SimSun", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, color="D9E2F3"):
    """表格单元格底纹"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_rich_para(doc, text, size=11, indent=False, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_cn_font(r, size=size, bold=True)
        else:
            r = p.add_run(part)
            set_cn_font(r, size=size)
    return p


def parse_table(lines):
    data_lines = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    if not data_lines:
        return [], []
    headers = [c.strip() for c in data_lines[0].strip("|").split("|")]
    rows = []
    for l in data_lines[1:]:
        cells = [c.strip() for c in l.strip("|").split("|")]
        cells += [""] * (len(headers) - len(cells))
        rows.append(cells[: len(headers)])
    return headers, rows


def build_docx(md_path, out_path, img_base):
    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    # Normal样式
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            title = m.group(2)
            h = doc.add_heading("", level=min(level, 4))
            h.paragraph_format.space_before = Pt(10 if level > 1 else 16)
            h.paragraph_format.space_after = Pt(6)
            if level == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = h.add_run(title)
                set_cn_font(r, size=18, bold=True)
            elif level == 2:
                r = h.add_run(title)
                set_cn_font(r, size=14, bold=True)
            elif level == 3:
                r = h.add_run(title)
                set_cn_font(r, size=12, bold=True)
            else:
                r = h.add_run(title)
                set_cn_font(r, size=11, bold=True)
            i += 1
            continue
        # 图片
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if m:
            cap = m.group(1)
            img_path = m.group(2)
            if not os.path.isabs(img_path):
                img_path = os.path.join(img_base, img_path)
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Cm(14))
                    # 图题
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(cap)
                    set_cn_font(cr, size=9, bold=True, color=RGBColor(0x44, 0x44, 0x44))
                except Exception as e:
                    add_rich_para(doc, f"[图缺失: {img_path}]", size=9)
            else:
                add_rich_para(doc, f"[图缺失: {img_path}]", size=9)
            i += 1
            continue
        # 表格
        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            headers, rows = parse_table(tbl_lines)
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                # 表头
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = ""
                    shade_cell(cell, "D9E2F3")
                    r = cell.paragraphs[0].add_run(h)
                    set_cn_font(r, size=9.5, bold=True)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 数据行
                for ri, row in enumerate(rows):
                    for j, val in enumerate(row):
                        cell = table.rows[ri + 1].cells[j]
                        cell.text = ""
                        parts = re.split(r"(\*\*.*?\*\*)", val)
                        for part in parts:
                            if not part:
                                continue
                            if part.startswith("**") and part.endswith("**"):
                                r = cell.paragraphs[0].add_run(part[2:-2])
                                set_cn_font(r, size=9.5, bold=True)
                            else:
                                r = cell.paragraphs[0].add_run(part)
                                set_cn_font(r, size=9.5)
                # 表后空行
                doc.add_paragraph()
            continue
        # 无序列表
        if line.startswith("- "):
            add_rich_para(doc, "· " + line[2:], size=11)
            i += 1
            continue
        # 普通段落
        add_rich_para(doc, line, size=11)
        i += 1

    doc.save(out_path)
    print(f"DOCX生成: {out_path}")


if __name__ == "__main__":
    md_path = sys.argv[1]
    out_path = sys.argv[2]
    img_base = sys.argv[3] if len(sys.argv) > 3 else ""
    build_docx(md_path, out_path, img_base)
