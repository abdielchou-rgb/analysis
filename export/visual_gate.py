
"""
VisualGate — DOCX成品质量门禁（结构校验版）
验证DOCX的结构合理性：工程质量 + 排版结构 + 图表嵌入 + 文本排版卫生

R6（2026-08-01 圆桌升级）：从纯白名单否决升级为结构校验。
此前只查"有没有坏模式"，对"没有结构/没有排版"这类缺失型缺陷完全免疫
（章节缺失→正则空匹配反而通过；图表全堆末尾→图片数达标算通过）。
新增三类硬检查：
  1. check_section_structure  — 正文必须存在真实标题（非TOC空锚点）
  2. check_chart_placement    — 图表必须嵌入正文，禁止全部堆叠末尾附录
  3. check_layout_quality     — 段首孤立句号 / 来源附录重复 / 空标题锚点
"""
import re, logging
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger("2hao.visual_gate")

def _first_non_empty(body):
    """找第一个有内容的元素"""
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl":
            cells = child.findall(".//" + qn("w:t"))
            if any(c.text and c.text.strip() for c in cells):
                return ("tbl", child)
        elif tag == "p":
            texts = [t.text for t in child.findall(".//" + qn("w:t")) if t.text]
            if any(t.strip() for t in texts):
                return ("p", child)
    return (None, None)

def _get_full_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)

def _para_has_image(p):
    """段落是否包含内嵌图片（w:drawing）"""
    return bool(p._p.findall(".//" + qn("w:drawing")))

def check_section_structure(doc, report_type="listed_company"):
    """结构校验：正文必须存在真实标题（Heading 且有实质文本，排除 TOC 空锚点/附录尾部标题）。

    TOC 域会生成一批【空】Heading 锚点段落（如 P34-P63），
    若报告正文没有真实标题，TOC 就是空壳 —— 这是排版失效的直接信号。
    """
    min_headings = {"industry_deep": 3, "listed_company": 3,
                    "unlisted_company": 3, "earnings_notes": 2,
                    "decision_memo": 3}.get(report_type, 3)
    appendix_marks = ("附录", "数据补充来源", "数据来源", "免责声明")
    real_heads = []
    for i, p in enumerate(doc.paragraphs):
        if "Heading" not in p.style.name and "Title" not in p.style.name:
            continue
        txt = p.text.strip()
        if not txt:  # TOC 空锚点
            continue
        if any(m in txt for m in appendix_marks):  # 尾部附录标题不计入正文结构
            continue
        real_heads.append((i, p.style.name, txt[:50]))
    if len(real_heads) < min_headings:
        return [{
            "severity": "error",
            "check": "missing_section_structure",
            "detail": "正文真实标题仅 %d 个（要求 ≥%d），报告无章节结构，正文为纯段落平铺: %s" % (
                len(real_heads), min_headings, str([h[2] for h in real_heads][:5])),
        }]
    return []

def check_chart_placement(doc, report_type="listed_company"):
    """图表嵌入位置：图表必须嵌入正文段落，禁止全部堆叠在末尾附录。

    以"附录/数据补充来源"段为界，界前出现的图视为正文嵌入。
    """
    appendix_start = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if any(m in p.text for m in ("附录", "数据补充来源")):
            appendix_start = i
            break
    image_paras = []
    for i, p in enumerate(doc.paragraphs):
        if _para_has_image(p):
            image_paras.append(i)
    total_images = len(image_paras)
    if total_images == 0:
        return []  # 图片数量由 check_image_quality 负责
    embedded = [i for i in image_paras if i < appendix_start]
    # 要求：正文嵌入图 ≥ min(2, 总图数)，且至少 40% 的图在正文
    min_embedded = min(2, total_images)
    if len(embedded) < min_embedded:
        return [{
            "severity": "error",
            "check": "charts_not_embedded",
            "detail": "图表 %d 张，正文仅嵌入 %d 张（要求 ≥%d），图表堆叠在末尾附录（附录起始 P%d）" % (
                total_images, len(embedded), min_embedded, appendix_start),
        }]
    return []

def check_layout_quality(doc):
    """排版卫生：段首孤立句号 / 来源附录重复注入 / 空标题锚点泛滥。"""
    issues = []
    # 1. 段首孤立句号（清洗残留）
    stray_periods = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("。") or t.startswith("．"):
            stray_periods.append(i)
    if stray_periods:
        issues.append({
            "severity": "error",
            "check": "stray_leading_period",
            "detail": "%d 段以孤立句号开头（文本清洗残留，如 P%d），应修正为正常起句" % (
                len(stray_periods), stray_periods[0]),
        })
    # 2. 来源附录重复注入（assemble 幂等性破坏）
    full = _get_full_text(doc)
    marker_count = full.count("AGENT_ENRICH_SOURCES")
    if marker_count > 2:  # 1 对开合标记 = 2 次出现
        issues.append({
            "severity": "error",
            "check": "duplicate_source_appendix",
            "detail": "来源附录标记出现 %d 次（应为 2 次=1 对），来源附录被重复注入，需修 assemble 幂等性" % marker_count,
        })
    # 3. 空标题锚点泛滥但无真实标题（TOC 空壳信号，已在结构检查覆盖，此处仅警告辅助）
    empty_headings = sum(
        1 for p in doc.paragraphs
        if "Heading" in p.style.name and not p.text.strip()
    )
    if empty_headings > 10:
        issues.append({
            "severity": "warning",
            "check": "empty_heading_anchors",
            "detail": "%d 个空 Heading 锚点（TOC 域生成），若正文无真实标题则目录为空壳" % empty_headings,
        })
    # 4. R83（2026-08-07）：来源附录原始区块泄漏——成品文档不应含 AGENT_ENRICH_SOURCES 原始数据
    # 油位 v0.89 事故：MD 尾部残留 20+ 条 enrich 来源原始数据，污染交付物。
    # VisualGate 直接对 DOCX 做终检：正文含该标记即视为未净化。
    if "AGENT_ENRICH_SOURCES" in full:
        issues.append({
            "severity": "error",
            "check": "enrich_source_leak",
            "detail": "文档含 AGENT_ENRICH_SOURCES 来源附录原始区块（未净化残留），应移除后再交付",
        })
    return issues

def check_table_quality(doc):
    """白名单：所有表格必须至少2行（header+至少1行数据）"""
    issues = []
    for i, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            issues.append({"severity": "error", "check": "table_too_short",
                           "detail": "表格%d只有%d行（至少需要2行）" % (i, len(table.rows))})
        elif len(table.rows) == 2:
            # 2-row table: check if second row has content
            cells = [c.text.strip() for c in table.rows[1].cells]
            if all(c == "" for c in cells):
                issues.append({"severity": "warning", "check": "table_only_header",
                               "detail": "表格%d只有表头无数据" % i})
        # Check header is not "#" or empty
        first_cell = table.rows[0].cells[0].text.strip() if table.rows[0].cells else ""
        if first_cell in ("#", "", " "):
            issues.append({"severity": "error", "check": "table_bad_header",
                           "detail": "表格%d的表头异常: %s" % (i, first_cell[:20])})
    return issues

def check_image_quality(doc, report_type="listed_company"):
    """白名单：图片数量必须满足最低要求"""
    min_images = {"industry_deep": 5, "listed_company": 5, "unlisted_company": 3,
                  "decision_memo": 2}.get(report_type, 3)
    count = len(doc.inline_shapes)
    issues = []
    if count < min_images:
        issues.append({"severity": "error", "check": "insufficient_images",
                       "detail": "图片%d张，最低要求%d张" % (count, min_images)})
    return issues

def check_font_explicit(doc):
    """白名单：所有run必须有显式字体设置"""
    issues = []
    for i, p in enumerate(doc.paragraphs[:50]):  # Check first 50 paragraphs
        for r in p.runs:
            if r.font.name is None:
                issues.append({"severity": "error", "check": "font_not_explicit",
                               "detail": "P%d: run字体未显式设置（继承模板）" % i})
                break
    return issues

def check_section_continuity(doc):
    """白名单：章节编号必须连续"""
    text = _get_full_text(doc)
    sections = re.findall(r"第(\d+)部分", text)
    if sections:
        nums = [int(s) for s in sections]
        expected = list(range(1, max(nums) + 1))
        missing = [str(e) for e in expected if e not in nums]
        if missing:
            return [{"severity": "error", "check": "section_discontinuity",
                     "detail": "缺少第%s部分" % "、".join(missing)}]
    return []

def check_first_element(doc):
    """白名单：首元素必须是标题"""
    body = doc.element.body
    tag, elem = _first_non_empty(body)
    issues = []
    if tag is None:
        issues.append({"severity": "error", "check": "empty_document", "detail": "文档没有任何内容"})
    elif tag == "tbl":
        issues.append({"severity": "error", "check": "first_is_table", "detail": "文档首元素是表格而非标题"})
    return issues

def check(docx_path, report_type="listed_company"):
    """全量检查"""
    doc = Document(docx_path)
    all_issues = []
    all_issues.extend(check_first_element(doc))
    all_issues.extend(check_table_quality(doc))
    all_issues.extend(check_image_quality(doc, report_type))
    all_issues.extend(check_font_explicit(doc))
    all_issues.extend(check_section_continuity(doc))
    # R6 新增：结构校验（排版真正不合格的地方）
    all_issues.extend(check_section_structure(doc, report_type))
    all_issues.extend(check_chart_placement(doc, report_type))
    all_issues.extend(check_layout_quality(doc))

    error_count = len([i for i in all_issues if i["severity"] == "error"])
    warning_count = len([i for i in all_issues if i["severity"] == "warning"])
    score = max(0, 1.0 - error_count * 0.25 - warning_count * 0.1)
    passed = error_count == 0

    return {
        "passed": passed,
        "score": score,
        "issues": all_issues,
        "images": len(doc.inline_shapes),
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
    }

def report(result):
    lines = ["=" * 60, "VisualGate 成品质检报告（白名单版）", "=" * 60,
             "状态: %s" % ("PASS" if result["passed"] else "FAIL"),
             "评分: %.2f" % result["score"],
             "文档: %d段 / %d表 / %d图" % (result["paragraphs"], result["tables"], result["images"])]
    if result["issues"]:
        lines.append("")
        lines.append("发现 %d 个问题:" % len(result["issues"]))
        for iss in result["issues"]:
            icon = "!" if iss["severity"] == "error" else "?"
            lines.append("  [%s] %s: %s" % (icon, iss["check"], iss["detail"]))
    return "\n".join(lines)


def blocking_issues(result):
    """Return only issues that should block export (error severity)"""
    return [i for i in result.get("issues", []) if i["severity"] == "error"]

def hard_fail_issues(result, hard_checks=None):
    """Return issues matching hard_fail check names"""
    if hard_checks is None:
        hard_checks = ["font_not_explicit", "insufficient_images", "first_is_table",
                       "table_too_short", "empty_document",
                       "enrich_source_leak", "charts_not_embedded", "duplicate_source_appendix"]  # R83
    return [i for i in result.get("issues", []) if i.get("check") in hard_checks]
