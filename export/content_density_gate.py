"""
Content Density Gate — 内容密度门禁

校验MD和DOCX的内容完整性：
- check_md(): 检查MD字数/图表/表格是否达标
- check_docx(): 检查DOCX字数/图片嵌入/表格
- reconcile(): 对比MD与DOCX字数，防内容丢失

所有方法返回: {"passed": bool, "score": float, "details": str, "warnings": [str]}
"""

import re
import sys
from pathlib import Path

_ANALYST_ROOT = Path(__file__).resolve().parent.parent
if str(_ANALYST_ROOT) not in sys.path:
    sys.path.insert(0, str(_ANALYST_ROOT))


class ContentDensityGate:
    """内容密度门禁——校验MD和DOCX的内容完整性"""

    MIN_CONTENT = {
        "industry_deep": {"chars": 8000, "charts": 5, "tables": 3},
        "listed_company": {"chars": 6000, "charts": 5, "tables": 3},
        "unlisted_company": {"chars": 5000, "charts": 4, "tables": 2},
        "earnings_notes": {"chars": 3000, "charts": 2, "tables": 1},
    }

    def __init__(self):
        pass

    # ----------------------------------------------------------------
    # MD checks
    # ----------------------------------------------------------------

    def check_md(self, md_text: str, report_type: str = "industry_deep") -> dict:
        """Check MD against minimum char/chart/table requirements.

        Returns dict with 'passed', 'score', 'details', 'warnings'.
        """
        thresholds = self.MIN_CONTENT.get(report_type, self.MIN_CONTENT["industry_deep"])
        warnings = []

        # Count plain-text chars (strip markdown syntax)
        plain = self._strip_markdown(md_text)

        # Count charts: ![alt](path)
        chart_count = len(re.findall(r"!\[.*?\]\(.*?\)", md_text))

        # Count tables: detectable markdown tables (header row + separator + data)
        tables = self._count_md_tables(md_text)

        chars_ok = len(plain) >= thresholds["chars"]
        charts_ok = chart_count >= thresholds["charts"]
        tables_ok = tables >= thresholds["tables"]

        all_ok = chars_ok and charts_ok and tables_ok

        if not chars_ok:
            warnings.append("MD字数不足: {} (需 >= {})".format(len(plain), thresholds["chars"]))
        if not charts_ok:
            warnings.append("MD图表不足: {}张 (需 >= {})".format(chart_count, thresholds["charts"]))
        if not tables_ok:
            warnings.append("MD表格不足: {}个 (需 >= {})".format(tables, thresholds["tables"]))

        # Score: weighted average of the three dimensions
        char_score = min(len(plain) / max(thresholds["chars"], 1), 1.0)
        chart_score = min(chart_count / max(thresholds["charts"], 1), 1.0)
        table_score = min(tables / max(thresholds["tables"], 1), 1.0)
        score = char_score * 0.4 + chart_score * 0.35 + table_score * 0.25

        detail_parts = [
            "字数: {}/{} ({})".format(len(plain), thresholds["chars"], "✓" if chars_ok else "✗"),
            "图表: {}/{} ({})".format(chart_count, thresholds["charts"], "✓" if charts_ok else "✗"),
            "表格: {}/{} ({})".format(tables, thresholds["tables"], "✓" if tables_ok else "✗"),
        ]

        return {
            "passed": all_ok,
            "score": round(score, 3),
            "details": " | ".join(detail_parts),
            "warnings": warnings,
            "char_count": len(plain),
            "chart_count": chart_count,
            "table_count": tables,
        }

    # ----------------------------------------------------------------
    # DOCX checks
    # ----------------------------------------------------------------

    def check_docx(self, docx_path: str, md_text: str | None = None, report_type: str = "industry_deep") -> dict:
        """Open DOCX, check char count / embedded image count / table count.

        If md_text is provided, also runs reconcile() internally.
        """
        thresholds = self.MIN_CONTENT.get(report_type, self.MIN_CONTENT["industry_deep"])
        warnings = []

        path = Path(docx_path)
        if not path.exists():
            return {
                "passed": False,
                "score": 0.0,
                "details": f"DOCX文件不存在: {docx_path}",
                "warnings": ["DOCX文件不存在"],
            }

        try:
            from docx import Document

            doc = Document(str(path))
        except Exception as e:
            return {
                "passed": False,
                "score": 0.0,
                "details": f"无法打开DOCX: {e}",
                "warnings": [f"无法打开DOCX: {e}"],
            }

        # Count characters from all paragraphs + table cells
        docx_chars = 0
        for para in doc.paragraphs:
            docx_chars += len(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_chars += len(cell.text)

        # Count embedded images (inline shapes)
        img_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if hasattr(run, "_element"):
                    drawings = run._element.findall(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    img_count += len(drawings)
        # Also check document-wide inline shapes
        for inline in doc.inline_shapes:
            if inline.type == 3:  # PICTURE
                img_count += 1

        # Deduplicate: inline_shapes already covers most cases
        # Use inline_shapes as primary source
        img_count = sum(1 for s in doc.inline_shapes if s.type == 3)

        table_count = len(doc.tables)

        chars_ok = docx_chars >= thresholds["chars"]
        imgs_ok = img_count >= thresholds["charts"]
        tables_ok = table_count >= thresholds["tables"]
        all_ok = chars_ok and imgs_ok and tables_ok

        if not chars_ok:
            warnings.append("DOCX字数不足: {} (需 >= {})".format(docx_chars, thresholds["chars"]))
        if not imgs_ok:
            warnings.append("DOCX图片不足: {}张 (需 >= {})".format(img_count, thresholds["charts"]))
        if not tables_ok:
            warnings.append("DOCX表格不足: {}个 (需 >= {})".format(table_count, thresholds["tables"]))

        char_score = min(docx_chars / max(thresholds["chars"], 1), 1.0)
        img_score = min(img_count / max(thresholds["charts"], 1), 1.0)
        table_score_d = min(table_count / max(thresholds["tables"], 1), 1.0)
        score = char_score * 0.4 + img_score * 0.35 + table_score_d * 0.25

        detail_parts = [
            "字数: {}/{} ({})".format(docx_chars, thresholds["chars"], "✓" if chars_ok else "✗"),
            "图片: {}/{} ({})".format(img_count, thresholds["charts"], "✓" if imgs_ok else "✗"),
            "表格: {}/{} ({})".format(table_count, thresholds["tables"], "✓" if tables_ok else "✗"),
        ]

        result = {
            "passed": all_ok,
            "score": round(score, 3),
            "details": " | ".join(detail_parts),
            "warnings": warnings,
            "char_count": docx_chars,
            "chart_count": img_count,
            "table_count": table_count,
        }

        # Reconcile if MD text provided
        if md_text:
            plain = self._strip_markdown(md_text)
            rec = self.reconcile(len(plain), docx_chars)
            if not rec["passed"]:
                result["passed"] = False
                result["warnings"].append(rec["details"])
                result["details"] += " | " + rec["details"]
                result["score"] = min(result["score"], rec["score"])

        return result

    # ----------------------------------------------------------------
    # Reconciliation
    # ----------------------------------------------------------------

    def reconcile(self, md_chars: int, docx_chars: int, threshold: float = 0.7) -> dict:
        """MD chars should be >= DOCX chars * threshold.

        If docx_chars < md_chars * threshold, content was lost during conversion.
        """
        if md_chars <= 0:
            return {"passed": True, "score": 1.0, "details": "MD为空，跳过对比", "warnings": []}

        if docx_chars <= 0:
            return {
                "passed": False,
                "score": 0.0,
                "details": f"DOCX无内容，转换可能失败 (MD: {md_chars} chars)",
                "warnings": ["DOCX无内容"],
            }

        ratio = docx_chars / md_chars
        passed = ratio >= threshold
        loss_pct = (1 - ratio) * 100 if ratio < 1 else 0

        if passed:
            score = min(ratio, 1.0)
            details = "MD/DOCX一致: {:.1%} (阈值 {:.0%}){}".format(ratio, threshold, " OK" if passed else "")
        else:
            # Score scales: 0.7 threshold = 0.0 score, 1.0 ratio = 1.0 score
            score = max(0.0, (ratio - threshold * 0.5) / (1 - threshold * 0.5))
            score = min(score, 1.0)
            details = f"内容丢失: DOCX={docx_chars} vs MD={md_chars} ({ratio:.1%}, 需 >= {threshold:.0%}), 丢失约{loss_pct:.0f}%"

        return {
            "passed": passed,
            "score": round(score, 3),
            "details": details,
            "warnings": [] if passed else [f"内容丢失: {loss_pct:.0f}%"],
            "loss_pct": round(loss_pct, 1),
            "ratio": round(ratio, 4),
        }

    # ----------------------------------------------------------------
    # Helpers
    # ----------------------------------------------------------------

    def _strip_markdown(self, text: str) -> str:
        """Strip markdown syntax to get approximate plain-text content."""
        # Remove code blocks
        text = re.sub(r"```[\s\S]*?```", "", text)
        # Remove images
        text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
        # Remove links, keep text
        text = re.sub(r"\[([^\]]*)\]\(.*?\)", r"\1", text)
        # Remove heading markers
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        # Remove bold/italic markers
        text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
        text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
        # Remove horizontal rules
        text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
        # Remove blockquote markers
        text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
        # Remove list markers
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
        # Remove inline code
        text = re.sub(r"`[^`]*`", "", text)
        return text

    def _count_md_tables(self, text: str) -> int:
        """Count markdown tables (header row + separator row required)."""
        lines = text.split("\n")
        count = 0
        i = 0
        while i < len(lines) - 1:
            line = lines[i].strip()
            if line.startswith("|") and line.endswith("|") and line.count("|") >= 3:
                # Check next line is a separator
                next_line = lines[i + 1].strip()
                if re.match(r"^\|[\s\-:]+\|", next_line):
                    count += 1
                    # Skip past this table
                    i += 2
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        i += 1
                    continue
            i += 1
        return count


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Content Density Gate")
    parser.add_argument("input_path", help="MD or DOCX file path")
    parser.add_argument(
        "--type",
        default="industry_deep",
        choices=["industry_deep", "listed_company", "unlisted_company", "earnings_notes"],
    )
    parser.add_argument("--md", help="Optional MD text for reconciliation")
    args = parser.parse_args()

    gate = ContentDensityGate()
    path = Path(args.input_path)

    if path.suffix.lower() == ".md":
        text = path.read_text(encoding="utf-8")
        result = gate.check_md(text, args.type)
    elif path.suffix.lower() == ".docx":
        md_text = None
        if args.md:
            md_text = Path(args.md).read_text(encoding="utf-8")
        result = gate.check_docx(str(path), md_text, args.type)
    else:
        print(f"Unsupported file type: {path.suffix}")
        return 1

    import json

    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["passed"] else 1


if __name__ == "__main__":
    sys.exit(main())
