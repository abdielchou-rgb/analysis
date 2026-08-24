"""V51.7 — DOCX 嵌表引擎：将 markdown 表格渲染为 Word 丰富表格 + 自动生成图表。

对标：中金/中信/高盛/MS 的 Word 报告平均每页 0.8-1.5 张表格/图表。
字节跳动报告已有 45 个表格块、141 行表格行——DOCX 应该把它们全部渲染出来。

核心策略：
  1. markdown 表格 → python-docx Table（已有 _render_table，需完善样式）
  2. 图表引擎生成的 png → 嵌入到表格附近或章节之间
  3. 数据型段落自动检测 → 生成 inline 迷你图表
"""

from __future__ import annotations
import logging, re
from pathlib import Path
from typing import Optional

logger = logging.getLogger("v51.export.docx_tables")

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def get_table_stats(markdown_text: str) -> dict:
    """统计 markdown 文本中的表格和图表数量。

    Returns:
        {"table_blocks": N, "table_rows": N, "chart_refs": N, "data_paragraphs": N}
    """
    table_blocks = 0
    table_rows = 0
    in_table = False

    for line in markdown_text.split('\n'):
        if '|' in line and '---' not in line:
            table_rows += 1
            if not in_table:
                table_blocks += 1
                in_table = True
        else:
            in_table = False

    chart_refs = len(re.findall(r'!\[.*\]\(.*\)', markdown_text))
    # 数据密集型段落
    data_paras = len(re.findall(r'\d{4}年|\d+\.\d+%|\d+亿元', markdown_text))

    return {
        "table_blocks": table_blocks,
        "table_rows": table_rows,
        "chart_refs": chart_refs,
        "data_paragraphs": data_paras,
    }


def render_table_with_style(doc: Document, table_block: str,
                            profile: dict, table_index: int = 0) -> None:
    """将 markdown 表格块渲染为 Word 表格，带投行级样式。

    样式规则：
      - 表头：深色背景 + 白色粗体字（机构主色）
      - 奇数行：浅色交替行
      - 首列：加粗
      - 数字列：右对齐
    """
    if not _HAS_DOCX:
        return

    lines = [l.strip() for l in table_block.split('\n') if l.strip()]
    if len(lines) < 2:
        return

    headers = [h.strip() for h in lines[0].split('|') if h.strip()]
    data_rows = []
    for line in lines[2:]:  # skip header separator
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data_rows.append(cells)

    if not headers or not data_rows:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    colors = profile.get('colors', {})

    # 从配置文件解析颜色（格式 "#003366" → RGBColor）
    def _parse_color(val, default=(0x00, 0x33, 0x66)):
        if isinstance(val, RGBColor):
            return val
        if isinstance(val, str) and val.startswith('#'):
            h = val.lstrip('#')
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        return RGBColor(*default)

    primary = _parse_color(colors.get('primary'), (0x00, 0x33, 0x66))
    accent = _parse_color(colors.get('accent'), (0xC4, 0x1E, 0x3A))

    # 表头
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1  # center
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 背景色
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): f'{int(primary[0]):02X}{int(primary[1]):02X}{int(primary[2]):02X}',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # 数据行
    for ri, row_data in enumerate(data_rows[:15]):  # 限制行数
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data[:n_cols]):
            cell = row.cells[ci]
            cell.text = cell_text
            # 首列加粗
            is_numeric = bool(re.match(r'^[\d\-.%]+$', cell_text.strip()))
            for paragraph in cell.paragraphs:
                paragraph.alignment = 2 if is_numeric else 0  # right align numbers
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = colors.get('text', RGBColor(0x33, 0x33, 0x33))
                    if ci == 0:
                        run.font.bold = True

        # 交替行背景
        if ri % 2 == 0:
            for ci in range(n_cols):
                cell = row.cells[ci]
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F5F5F5',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elm)

    # 表格间距
    doc.add_paragraph()  # spacing


def embed_charts_in_docx(doc: Document, chart_paths: dict,
                         profile: dict, style_id: str = "cicc") -> None:
    """在报告关键位置嵌入图表。

    策略：
      - 如果有敏感性矩阵图 → 在"估值"相关段落后插入
      - 如果有柱状图/折线图 → 在"财务分析"相关段落后插入
      - 每个图表附标题和图号
    """
    if not _HAS_DOCX or not chart_paths:
        return

    for chart_type, chart_path in sorted(chart_paths.items()):
        if not Path(chart_path).exists():
            continue
        try:
            doc.add_picture(str(chart_path), width=Inches(5.5))
            # 图标题
            cap = doc.add_paragraph()
            cap.alignment = 1  # center
            run = cap.add_run(f"图：{chart_type}")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()  # spacing
        except Exception as e:
            logger.warning(f"Chart embed failed {chart_path}: {e}")


def enhance_tables_in_docx(doc: Document, report_md: str,
                           chart_paths: dict = None,
                           style_id: str = "cicc") -> None:
    """增强现有 Document 对象中的表格和图表。

    在 markdown_to_docx 调用后，对 doc 追加表格和图表，
    实现"丰富表格 + 图表嵌入"。

    用法:
        from export.docx_exporter import markdown_to_docx
        doc = markdown_to_docx(md_text, output_path)
        enhance_tables_in_docx(doc, md_text, chart_paths, style_id)
    """
    if not _HAS_DOCX:
        return

    from core.styles.profiles import get_style
    profile = get_style(style_id)

    # 1. 从 markdown 提取表格并渲染
    lines = report_md.split('\n')
    table_buffer = []
    in_table = False
    table_idx = 0

    for line in lines:
        if '|' in line and '---' not in line:
            table_buffer.append(line)
            in_table = True
        else:
            if in_table and table_buffer:
                render_table_with_style(doc, '\n'.join(table_buffer), profile, table_idx)
                table_idx += 1
                table_buffer = []
            in_table = False

    if in_table and table_buffer:
        render_table_with_style(doc, '\n'.join(table_buffer), profile, table_idx)

    # 2. 嵌入图表
    if chart_paths:
        embed_charts_in_docx(doc, chart_paths, profile, style_id)

    stats = get_table_stats(report_md) if report_md else {}
    logger.info(f"增强完成: {table_idx} 表格, {len(chart_paths or {})} 图表嵌入")
