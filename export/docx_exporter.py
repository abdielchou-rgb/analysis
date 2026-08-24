#!/usr/bin/env python3
"""docx_exporter.py — Template-based DOCX rendering (Phase 2).
Uses cicc.dotx template with predefined styles instead of programmatic generation.
"""

import logging
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("2hao.docx_exporter")

ROOT = Path(__file__).resolve().parent.parent


class TemplateDocxExporter:
    """DOCX exporter using .dotx template with predefined styles."""

    STYLE_MAP = {
        "h1": "CICCH1",
        "h2": "CICCH2",
        "h3": "CICCH2",
        "body": "CICCBody",
        "title": "CICCTitle",
        "takeaway": "CICCKeyTakeaway",
    }

    def __init__(self, style_id="cicc"):
        self.style_id = style_id
        template_path = ROOT / "templates" / (style_id + ".dotx")
        if template_path.exists():
            self.template = str(template_path)
            logger.info("Using template: %s", self.template)
        else:
            self.template = None
            logger.warning("Template not found: %s, using empty", template_path)

    def export(self, md_text, output_path, title="", subtitle="", author="2hao Analyst"):
        """Convert markdown to DOCX using template."""
        if self.template:
            doc = Document(self.template)
        else:
            doc = Document()

        # Clear sample content from template
        for p in doc.paragraphs:
            p.clear()
        # Remove sample tables
        for tbl in doc.tables:
            tbl._element.getparent().remove(tbl._element)

        # Parse markdown into sections
        lines = md_text.split("\n")
        i = 0
        in_table = False
        table_rows = []
        header_row = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines and separators
            if not stripped or stripped.startswith("---"):
                if in_table:
                    self._add_table(doc, header_row, table_rows)
                    in_table = False
                    table_rows = []
                    header_row = []
                i += 1
                continue

            # Handle tables
            if "|" in stripped and "---" in stripped:
                # Table separator row - skip
                i += 1
                continue
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                if not in_table:
                    header_row = cells
                    in_table = True
                else:
                    table_rows.append(cells)
                i += 1
                continue

            # Flush any open table
            if in_table:
                self._add_table(doc, header_row, table_rows)
                in_table = False
                table_rows = []
                header_row = []

            # Headings
            if stripped.startswith("# "):
                self._add_paragraph(doc, stripped[2:].strip(), "title")
            elif stripped.startswith("## "):
                self._add_paragraph(doc, stripped[3:].strip(), "h1")
            elif stripped.startswith("### "):
                self._add_paragraph(doc, stripped[4:].strip(), "h2")

            # Bold takeaways
            elif stripped.startswith("**") and stripped.endswith("**"):
                self._add_paragraph(doc, stripped.strip("*"), "takeaway")

            # Images - skip (handled by content_placer)
            elif stripped.startswith("!["):
                # Image reference - try to embed from path
                img_match = re.match(r"!\[.*\]\((.+)\)", stripped)
                if img_match:
                    img_path = img_match.group(1)
                    self._add_image(doc, img_path)

            # Regular paragraphs
            elif len(stripped) > 10:
                self._add_paragraph(doc, stripped, "body")

            # Short lines as body
            elif stripped:
                self._add_paragraph(doc, stripped, "body")

            i += 1

        # Final flush
        if in_table:
            self._add_table(doc, header_row, table_rows)

        # Set document title
        if title:
            # Update first paragraph if it's a title
            for p in doc.paragraphs:
                if p.style and "Title" in str(p.style.name):
                    p.text = title
                    break

        # Save
        output_path = str(output_path)
        doc.save(output_path)
        logger.info("DOCX saved: %s (%d KB)", output_path, os.path.getsize(output_path) // 1024)
        return output_path

    def _add_paragraph(self, doc, text, style_key="body"):
        """Add paragraph with mapped style."""
        style_name = self.STYLE_MAP.get(style_key, "CICCBody")
        p = doc.add_paragraph(text, style=style_name)

        # Clean up any bold markers in text
        for run in p.runs:
            run.text = run.text.replace("**", "")

        return p

    def _add_table(self, doc, headers, rows):
        """Add formatted table."""
        if not headers and not rows:
            return
        ncols = max(len(headers), max((len(r) for r in rows), default=0))
        if ncols == 0:
            return

        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = "Table Grid"

        # Headers
        for i, h in enumerate(headers):
            if i < ncols:
                cell = table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                        run.font.name = "Arial"

        # Data rows
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci < ncols:
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = val
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

        doc.add_paragraph()  # spacing after table

    def _add_image(self, doc, img_path):
        """Add image to document, resizing to fit page width."""
        if not os.path.isfile(img_path):
            logger.debug("Image not found: %s", img_path)
            return
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            # Center the image
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add caption
            caption = doc.add_paragraph(os.path.basename(img_path), style="CICCBody")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        except Exception as e:
            logger.debug("Image embed failed: %s", e)


# Keep backward compatibility
CICCStylePDFExporter = TemplateDocxExporter  # alias for code that imports this

if __name__ == "__main__":
    exporter = TemplateDocxExporter("cicc")
    test_md = "# Test Report\n\n## Section 1\n\nThis is body text.\n\n| Col1 | Col2 |\n|------|------|\n| A | B |\n"
    exporter.export(test_md, "test_output.docx", title="Test Report")
    print("Test DOCX generated")


def export_docx(report_md, output_path, title=None, subtitle=None, author=None):
    exporter = TemplateDocxExporter()
    return exporter.export(
        report_md, output_path, title=title or "", subtitle=subtitle or "", author=author or "2hao Analyst"
    )


def markdown_to_docx(md_text, output_path, title="", subtitle="", author="2hao Analyst"):
    """R39：markdown → docx（复用 TemplateDocxExporter）。

    integrated_exporter 长期 import 此函数但从未定义 → docx 导出链路静默失败。
    现在补上定义，让导出真正工作。
    """
    exporter = TemplateDocxExporter()
    return exporter.export(md_text, output_path, title=title, subtitle=subtitle, author=author)


def add_static_toc(docx_path: str, md_text: str) -> int:
    """R42：向 docx 插入静态目录（从 markdown 标题提取，无需 Word 刷新）。

    背景：DOCX 目录为空问题——TOC 域需要 Word 打开后 F9 刷新才渲染，
    python-docx 无法生成真实目录内容。方案：解析 markdown 的 #/## 标题，
    生成静态目录段落插入正文前（封面后），读者一眼可见章节结构。
    返回插入的目录条目数。
    """
    from pathlib import Path as _P

    from docx import Document as _Doc
    from docx.shared import Pt as _Pt

    if not md_text or not _P(docx_path).exists():
        return 0
    # 提取 markdown 标题（兼容管线标题规范："# 报告主标题 + # 一级章节 + ## 小节 + ### 三级"）
    # R43（2026-08-02）：补单井号一级章节提取——柯力报告用 "# 一、公司概况" 作为顶级章节，
    # 上一版只取 ##/### 导致顶级章节缺失。跳过第一个 # 主标题（通常是报告名），其余 # 作一级章节。
    headings = []
    _seen_main_title = False
    for line in md_text.split("\n"):
        s = line.strip()
        if s.startswith("### "):
            headings.append((2, s[4:].strip()))
        elif s.startswith("## "):
            headings.append((1, s[3:].strip()))
        elif s.startswith("# "):
            title_text = s[2:].strip()
            if not _seen_main_title:
                # 第一个 # 视为报告主标题（跳过，不进目录）
                _seen_main_title = True
            else:
                headings.append((0, title_text))
    if not headings:
        return 0
    try:
        doc = _Doc(docx_path)
        # 找到封面标题段落（第一个非空段），在其后插入目录
        insert_idx = 0
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                insert_idx = i + 1
                break
        # 用 XML 插入：在封面段落后追加"目录标题 + 条目"（保持顺序）
        anchor = doc.paragraphs[insert_idx - 1]._p
        # 目录标题
        toc_title = doc.add_paragraph("目  录")
        try:
            toc_title.style = doc.styles["Heading 1"]
        except Exception:
            pass
        toc_title.alignment = 1  # center
        anchor.addnext(toc_title._p)
        # 条目（正序插入，用 addnext 维持顺序）
        prev_el = toc_title._p
        for level, text in headings:
            indent = "    " * level
            p = doc.add_paragraph(indent + text)
            run = p.runs[0] if p.runs else p.add_run(indent + text)
            run.font.size = _Pt(12 if level == 0 else 10 if level == 1 else 9)
            prev_el.addnext(p._p)
            prev_el = p._p
        doc.save(docx_path)
        return len(headings)
    except Exception as _e:
        logger.debug("[TOC] add_static_toc failed: %s", _e)
        return 0


def pandoc_to_docx(markdown_path, output_path, style="cicc"):
    """R39：优先尝试 pandoc 转换，失败回退 markdown_to_docx。

    返回生成的 docx 路径；失败返回 None（由调用方回退）。
    """
    from pathlib import Path as _P

    md_p = _P(markdown_path)
    out_p = _P(output_path)
    if not md_p.exists():
        return None
    try:
        # 尝试 pandoc（若系统已安装）
        import subprocess

        r = subprocess.run(
            ["pandoc", str(md_p), "-o", str(out_p)],
            capture_output=True,
            timeout=60,
        )
        if r.returncode == 0 and out_p.exists():
            return str(out_p)
    except Exception as _e:
        logger.debug("[PANDOC] not available, fallback: %s", _e)
    # 回退：python-docx 原生转换
    md_text = md_p.read_text(encoding="utf-8")
    try:
        markdown_to_docx(md_text, str(out_p), title=md_p.stem)
        if out_p.exists():
            return str(out_p)
    except Exception as _e:
        logger.warning("[PANDOC] fallback failed: %s", _e)
    return None


def clean_empty_paragraphs(docx_path: str) -> int:
    """R30 模块3（排版修复）：删除 docx 中连续的空段落（治空白页老问题）。

    Pandoc 转换会保留 markdown 连续空行 → docx 产生空段落 → 空白页。
    此函数 post-process：删除空 <w:p> 段落（保留含图/表格的段落）。

    返回删除的空段落数。
    """
    import re as _re
    import shutil
    import zipfile
    from pathlib import Path as _P

    p = _P(docx_path)
    if not p.exists():
        return 0
    tmp = p.with_suffix(".tmp.docx")
    removed = 0
    try:
        with zipfile.ZipFile(p, "r") as zin:
            names = zin.namelist()
            if "word/document.xml" not in names:
                return 0
            doc_xml = zin.read("word/document.xml").decode("utf-8")

            # 删除空段落：<w:p ...>...</w:p> 且无 <w:t>、无 <w:drawing>、无 <w:tbl>
            def _is_empty_para(m):
                nonlocal removed
                body = m.group(1)
                if "<w:t" in body or "<w:drawing>" in body or "<w:tbl>" in body:
                    return m.group(0)
                # 去掉标签后无文本 → 空段
                text = _re.sub(r"<[^>]+>", "", body).strip()
                if not text:
                    removed += 1
                    return ""
                return m.group(0)

            doc_xml = _re.sub(r"<w:p\b[^>]*>(.*?)</w:p>", _is_empty_para, doc_xml, flags=_re.S)
            # 写回
            with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        data = doc_xml.encode("utf-8")
                    zout.writestr(item, data)
        shutil.move(str(tmp), str(p))
    except Exception:
        try:
            if tmp.exists():
                tmp.unlink()
        except Exception:
            pass
        return 0
    return removed


def _scrub_aigc_artifacts(md_text: str) -> str:
    """R38（2026-08-02）：清理管线内部残留物（此前 integrated_exporter 引用
    此函数但从未定义，导致 AGENT_ENRICH_SOURCES 注释/内部字段名/病句泄漏进成品）。

    清理项：
      1. AGENT_ENRICH_SOURCES HTML 注释块
      2. "公司研究素材.xxx" 内部字段名引用
      3. 标点病句（"。。"、"。，"粘连）
      4. 括号计数残留（"共6个环节（6）"）
    """
    import re as _re

    text = md_text or ""
    # 1. AGENT_ENRICH_SOURCES 注释块（含内容）
    text = _re.sub(r"<!--\s*AGENT_ENRICH_SOURCES.*?/AGENT_ENRICH_SOURCES\s*-->", "", text, flags=_re.S)
    text = _re.sub(r"<!--\s*/?AGENT_ENRICH_SOURCES\s*-->", "", text)
    # 2. 内部字段名（公司研究素材.xxx）——连同前后连接词一起清除
    text = _re.sub(r"[^。；\n]*公司研究素材\.[A-Za-z_0-9]+[^。；\n]*[。；]?", "", text)
    text = _re.sub(r"[^。；\n]*公司研究素材\.[A-Za-z_0-9]+", "", text)
    # 3. 标点病句
    text = text.replace("。。", "。").replace("。，", "。")
    text = text.replace("，，", "，").replace("。。", "。")
    # 4. 括号计数残留（如"共6个环节（6）"、"驱动因素计数为5项（5）"）
    text = _re.sub(r"[（(](\d{1,3})[)）](?=\s*[。；\n]|$)", "", text)
    text = _re.sub(r"(共|计)\d{1,3}(个|项|条|种)(环节|驱动|因素|指标)[（(]?\d{1,3}[)）]?", r"\1\2\3", text)
    # 5. 空行压缩（清理后可能产生多余空行）
    text = _re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
