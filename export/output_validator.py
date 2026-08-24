"""V56 OutputValidator — 最终产出物格式校验器

校验 DOCX / PDF 文件的格式质量，确保没有 Markdown 工件，
字体一致，图表数量达标。

集成到报告生成 loop 中，作为评分→校准→输出 的最后一道门禁。
"""

from __future__ import annotations

import logging
import re
from typing import Any

logger = logging.getLogger("v56.validation.output")


class OutputValidator:
    """最终产出物格式校验器

    校验维度:
    1. 无 MD 工件（** __ ``` >）
    2. 图表数量达标
    3. 字体一致性
    4. 字号层级合理
    5. 页面布局规范
    6. 无 AIGC 元数据泄漏（人格头 / 模板桩）
    """

    # AIGC 人格头模式
    _AIGC_PERSONA_PATTERNS = [
        r"我是资深(行业)?分析师",
        r"以下是为您撰写的",
        r"为您撰写了",
    ]

    # 模板占位桩模式
    _TEMPLATE_PLACEHOLDER_PATTERNS = [
        r"在此处阐述",
        r"在此处填写",
        r"在此输入",
        r"\[报告标题\]",
        r"\[报告副标题\]",
    ]

    def __init__(self):
        self.results: dict[str, Any] = {}

    def validate_docx(self, filepath: str) -> dict[str, Any]:
        """校验 DOCX 文件格式"""
        try:
            from docx import Document

            doc = Document(filepath)
        except Exception as e:
            return {"error": f"Cannot open DOCX: {e}", "passed": False}

        results = {"format": "docx", "passed": True, "checks": {}}

        # 1. MD 工件检测
        md_artifacts = 0
        md_patterns = []
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text
                if "**" in text:
                    md_artifacts += text.count("**")
                    md_patterns.append(text[:50])
                if "__" in text:
                    md_artifacts += text.count("__")
                if "```" in text:
                    md_artifacts += text.count("```")

        results["checks"]["md_artifacts"] = {
            "count": md_artifacts,
            "passed": md_artifacts == 0,
            "samples": md_patterns[:3],
        }
        if md_artifacts > 0:
            results["passed"] = False

        # 2. 图表数量
        try:
            chart_count = len(doc.inline_shapes)
        except Exception:
            chart_count = 0
        results["checks"]["chart_count"] = {
            "count": chart_count,
            "minimum": 5,
            "passed": chart_count >= 5,
        }
        if chart_count < 3:
            results["passed"] = False

        # 3. 字体一致性
        fonts_seen = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_seen.add(run.font.name)
        results["checks"]["font_consistency"] = {
            "fonts": list(fonts_seen),
            "count": len(fonts_seen),
            "passed": len(fonts_seen) <= 3,  # max 3 fonts (body, heading, mono)
        }

        # 4. 总段落数和字数
        total_paras = len(doc.paragraphs)
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        results["checks"]["content_volume"] = {
            "paragraphs": total_paras,
            "characters": total_chars,
            "passed": total_chars >= 3000,
        }

        # 5. 图片位置检测（检查是否有图片在文档末尾而非正文中）
        last_chart_pos = 0
        first_chart_pos = 0
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith("[图表"):
                if first_chart_pos == 0:
                    first_chart_pos = idx
                last_chart_pos = idx

        results["checks"]["chart_position"] = {
            "first_chart_para": first_chart_pos,
            "last_chart_para": last_chart_pos,
            "passed": last_chart_pos < total_paras - 3,  # 图表不在最后3段之后
        }

        # 6. AIGC 元数据泄漏检测
        aigc_leaks = []
        template_leaks = []
        for para in doc.paragraphs:
            ptext = para.text
            for pattern in self._AIGC_PERSONA_PATTERNS:
                if re.search(pattern, ptext):
                    aigc_leaks.append(ptext[:80].strip())
            for pattern in self._TEMPLATE_PLACEHOLDER_PATTERNS:
                if re.search(pattern, ptext):
                    template_leaks.append(ptext[:80].strip())

        results["checks"]["aigc_leak"] = {
            "count": len(aigc_leaks),
            "passed": len(aigc_leaks) == 0,
            "samples": list(set(aigc_leaks))[:3],
        }
        results["checks"]["template_placeholder"] = {
            "count": len(template_leaks),
            "passed": len(template_leaks) == 0,
            "samples": list(set(template_leaks))[:3],
        }
        if aigc_leaks or template_leaks:
            results["passed"] = False

        self.results = results
        return results

    def validate_pdf(self, filepath: str) -> dict[str, Any]:
        """校验 PDF 文件格式"""
        results = {"format": "pdf", "passed": True, "checks": {}}
        try:
            # 使用 PyPDF2 或 pdfminer 检查
            try:
                import PyPDF2

                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                    results["checks"]["page_count"] = {
                        "count": page_count,
                        "passed": page_count >= 5,
                    }
            except ImportError:
                # 降级：检查文件大小
                import os

                size = os.path.getsize(filepath)
                results["checks"]["file_size"] = {
                    "bytes": size,
                    "passed": size > 50000,
                }
        except Exception as e:
            results["error"] = str(e)
            results["passed"] = False

        self.results = results
        return results

    def summary(self) -> str:
        """生成人类可读的校验报告"""
        if not self.results:
            return "未运行校验"

        lines = [f"格式校验报告 ({self.results.get('format', '?')})"]
        lines.append(f"总体: {'✅ 通过' if self.results.get('passed') else '❌ 失败'}")

        for check_name, check_data in self.results.get("checks", {}).items():
            status = "✅" if check_data.get("passed") else "❌"
            lines.append(f"  {status} {check_name}: {check_data}")

        return "\n".join(lines)


# 用于 report_calibrator.py 集成的快捷函数
def validate_report(filepath: str) -> dict:
    """校验报告文件格式的快捷函数"""
    validator = OutputValidator()
    if filepath.endswith(".docx"):
        return validator.validate_docx(filepath)
    elif filepath.endswith(".pdf"):
        return validator.validate_pdf(filepath)
    else:
        return {"error": f"Unsupported format: {filepath}", "passed": False}
