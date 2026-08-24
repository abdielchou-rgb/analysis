"""
2号分析师 — 报告导出器

支持导出为 Word (.docx) 和 PDF 格式。

Word 导出 (python-docx):
  - 保留 Markdown 标题层级 (H1-H4 → Word Heading 1-4)
  - 中文字体支持 (Noto/SimSun fallback)
  - 插入图表图片
  - 保留表格格式
  - 页眉页脚 (报告标题 · 页码)

PDF 导出:
  - 优先通过 python-docx → LibreOffice (subprocess) 转换
  - Fallback 到 reportlab 直接生成

Usage:
    from tools.exporter import ReportExporter
    exporter = ReportExporter()
    result = exporter.export_all(markdown_text, "output_name", chart_paths)
"""

from __future__ import annotations

import logging
import os
import re
import subprocess
import sys
from pathlib import Path

logger = logging.getLogger("v30.exporter")

# python-docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


class ReportExporter:
    """
    报告导出器：Markdown → Word/PDF。
    """

    def __init__(self, company_name: str = "", style_id: str = "cicc", title: str = ""):
        self.company_name = company_name
        self.style_id = style_id
        self.title = title
        self._doc: Document | None = None

    # ═══════════════════════════════════════════════
    # Word 导出
    # ═══════════════════════════════════════════════

    def to_docx(
        self,
        markdown_text: str,
        output_path: str,
        chart_paths: dict[str, str] | None = None,
    ) -> str:
        """
        将 Markdown 文本转换为 Word (.docx) 文档。

        Args:
            markdown_text: Markdown 格式的报告文本
            output_path: 输出 .docx 文件路径
        from export.pre_export_sheriff import sanitize, validate_sanitized
        markdown_text = sanitize(markdown_text)
            chart_paths: 图表路径字典，如 {"revenue_trend": "path/to/chart.png"}
                         会在 Markdown 中查找 ![](chart:key) 占位符并插入图片

        Returns:
            str: 生成的 .docx 文件路径
        """
        # Load from template if available
        template_path = Path(__file__).resolve().parent.parent / "templates" / self.style_id / "report.dotx"
        if template_path.exists():
            self._doc = Document(str(template_path))
            # FP4: 自动清理模板占位内容(段落+表格)
            for _p in self._doc.paragraphs:
                _p.clear()
            # 删除模板自带空表格
            for _t in list(self._doc.tables):
                _all_empty = all(c.text.strip() == "" for row in _t.rows for c in row.cells)
                if _all_empty or (_t.rows and _t.rows[0].cells[0].text.strip() in ("指标", "项目", "指标/项目")):
                    _t._element.getparent().remove(_t._element)
            logger.info("Loaded template: %s (placeholder cleared)", template_path)
        else:
            self._doc = Document()
        self._setup_styles()
        # Always add cover page and TOC
        self._add_cover_page(
            title_text=self.title if self.title else self.company_name,
            style_name=self.style_id.upper() if self.style_id else "CICC",
        )
        self._add_toc()
        self._setup_headers_footers()

        chart_paths = chart_paths or {}

        lines = markdown_text.split("\n")
        i = 0
        in_table = False
        table_rows: list[list[str]] = []
        table_col_count = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # ── 表格处理 ──
            if "|" in stripped and stripped.startswith("|") and stripped.endswith("|"):
                # 检测分隔行
                if re.match(r"^\|[\s\-:|]+\|$", stripped):
                    in_table = True
                    i += 1
                    continue

                if in_table or self._is_table_row(stripped):
                    in_table = True
                    cells = self._parse_table_row(stripped)
                    table_rows.append(cells)
                    table_col_count = max(table_col_count, len(cells))
                    i += 1
                    continue
            else:
                # 表格结束 — 渲染
                if in_table and table_rows:
                    self._render_table(table_rows)
                    table_rows = []
                    table_col_count = 0
                    in_table = False

            # ── 图表占位符 ──
            # 支持[CHART:fig_id]占位符(转换为![](chart:fig_id))
            stripped = re.sub(r"\{?CHART:(\w+)\\}?", r"![](chart:\1)", stripped)
            chart_match = re.match(r"!\[\]\(chart:(\w+)\)", stripped)
            if chart_match:
                chart_key = chart_match.group(1)
                if chart_key in chart_paths:
                    self._insert_chart(chart_paths[chart_key])
                else:
                    self._doc.add_paragraph(f"[图表未找到: {chart_key}]")
                i += 1
                continue

            # ── 图片 (一般路径) ──
            img_match = re.match(r"!\[.*?\]\((.+?)\)", stripped)
            if img_match:
                img_path = img_match.group(1)
                resolved = None
                # 1. Direct path
                if Path(img_path).exists():
                    resolved = img_path
                # 2. In chart_paths dict
                elif chart_paths:
                    for key, path in chart_paths.items():
                        if key in img_path or img_path in path:
                            resolved = path
                            break
                # 3. Search in output directories
                if not resolved:
                    p = Path(img_path)
                    for root in [Path("output"), Path("output/charts"), Path(".")]:
                        cand = root / p.name
                        if cand.exists():
                            resolved = str(cand)
                            break
                # 4. Absolute path
                if not resolved and img_path.startswith("D:"):
                    p = Path(img_path)
                    if p.exists():
                        resolved = str(p)
                # Insert
                if resolved and Path(resolved).exists():
                    self._insert_chart(resolved)
                    logger.info("  Embedded image: %s", resolved)
                else:
                    self._doc.add_paragraph("[Image: %s]" % img_path)
                    logger.warning("  Image not found: %s", img_path)
                i += 1
                continue

            # ── 水平线 ──
            if re.match(r"^---+\s*$", stripped) or re.match(r"^\*\*\*+\s*$", stripped):
                self._doc.add_paragraph("_" * 60)
                i += 1
                continue

            # ── 引用块 ──
            if stripped.startswith(">"):
                quote_text = stripped.lstrip("> ").strip()
                p = self._doc.add_paragraph()
                run = p.add_run(quote_text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.paragraph_format.left_indent = Cm(1)
                i += 1
                continue

            # ── 代码块 ──
            if stripped.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = self._doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.size = Pt(9)
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i += 1
                continue

            # ── 空行 ──
            if not stripped:
                # 除非前面是标题/表格, 否则跳过空行避免过多间距
                i += 1
                continue

            # ── 标题 (H1-H4) ──
            heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Remove bold markers from headings
                text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
                self._doc.add_heading(text, level=level)
                i += 1
                continue

            # ── 普通段落 ──
            p = self._doc.add_paragraph()
            self._add_formatted_text(p, stripped)
            i += 1

        # 最后的表格
        if in_table and table_rows:
            self._render_table(table_rows)

        # 强制所有 run 显式设置字体 (covering ALL document parts)
        all_runs = []
        # 正文段落
        for p in self._doc.paragraphs:
            all_runs.extend(p.runs)
        # 表格
        for t in self._doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        all_runs.extend(p.runs)
        # 页眉
        for section in self._doc.sections:
            for p in section.header.paragraphs:
                all_runs.extend(p.runs)
            for p in section.footer.paragraphs:
                all_runs.extend(p.runs)
        for run in all_runs:
            if run.font.name is None:
                try:
                    self._apply_run_font(run)
                except Exception:
                    pass

        # 清理连续空段落（保留封面所需的最少留白，表格后间隔压缩为1行）
        body = self._doc.element.body
        paras = body.findall(qn("w:p"))
        prev_empty = False
        cover_zone = True  # 封面区域（第一个分页符前）保留原样
        for p_el in paras:
            # 检测分页符，越过封面区域
            if cover_zone and p_el.findall(qn("w:r") + "/" + qn("w:br")):
                cover_zone = False
            if cover_zone:
                continue
            text = "".join(node.text or "" for node in p_el.iter() if node.tag == qn("w:t"))
            # 图片段落（含 drawing）不视为空段，避免连续图片中第二张被误删
            has_drawing = p_el.findall(".//" + qn("w:drawing")) or p_el.findall(".//" + qn("w:pict"))
            is_empty = (not text.strip()) and not has_drawing
            if is_empty and prev_empty:
                p_el.getparent().remove(p_el)
            else:
                prev_empty = is_empty

        self._doc.save(output_path)
        logger.info(f"Word 文档已保存: {output_path}")
        return output_path

    def _apply_run_font(self, run, style="cicc"):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        run.font.name = "SimHei"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "SimHei")

    def _add_formatted_text(self, paragraph, text: str):
        """解析行内格式（粗体、斜体）。"""
        # Split by bold (**text**), italic (*text*), and inline code (`text`)
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            else:
                paragraph.add_run(part)

    def _is_table_row(self, line: str) -> bool:
        return bool(re.match(r"^\|.+\|$", line))

    def _parse_table_row(self, line: str) -> list[str]:
        """解析 Markdown 表格行。"""
        cells = line.strip().split("|")
        # Remove empty first/last from leading/trailing |
        if cells and cells[0].strip() == "":
            cells = cells[1:]
        if cells and cells[-1].strip() == "":
            cells = cells[:-1]
        return [c.strip() for c in cells]

    def _render_table(self, rows: list[list[str]]):
        """渲染表格到 Word 文档。"""
        if not rows:
            return

        n_cols = max(len(r) for r in rows)
        n_rows = len(rows)

        table = self._doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for i, row_cells in enumerate(rows):
            for j in range(n_cols):
                cell = table.cell(i, j)
                text = row_cells[j] if j < len(row_cells) else ""
                # 去除行内粗体
                text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
                cell.text = text

                # 格式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "SimSun"

                # 表头加粗/背景色
                if i == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True

        # 不再添加表格后空段（由全局空段清理统一控制间距，避免空白页）

    def _insert_chart(self, chart_path: str):
        """在文档中插入图表图片（自适应页面宽度）。"""
        if not Path(chart_path).exists():
            self._doc.add_paragraph(f"[图表文件未找到: {chart_path}]")
            return

        try:
            display_w = 5.5
            try:
                from PIL import Image

                img = Image.open(chart_path)
                img_w = img.size[0]
                page_w = 6.3  # A4 usable width (inches)
                scale = min(page_w / (img_w / 96.0), 1.0) if img_w > 0 else 1.0
                display_w = min(img_w * scale / 96.0, page_w)
            except ImportError:
                pass

            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_path, width=Inches(display_w))
            # 图片后不再添加空段，由全局空段清理统一控制间距
        except Exception as e:
            logger.warning("Failed to insert chart: %s", e)
            self._doc.add_paragraph("[Image: %s]" % chart_path)

    def _setup_styles(self):
        """设置文档默认样式。"""
        style = self._doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        # 设置中文字体 fallback
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="SimSun"/>')
            rpr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), "SimSun")

        # 标题样式
        for level in range(1, 5):
            heading_style = self._doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            heading_style.font.bold = True
            if level == 1:
                heading_style.font.size = Pt(18)
            elif level == 2:
                heading_style.font.size = Pt(15)
            elif level == 3:
                heading_style.font.size = Pt(13)
            else:
                heading_style.font.size = Pt(11)

    def _setup_headers_footers(self):
        """设置页眉页脚。"""
        for section in self._doc.sections:
            # 页眉
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr = self.title if self.title else (self.company_name if self.company_name else "深度研究报告")
            run = hp.add_run(hdr)
            self._apply_run_font(run)

            # 页脚（页码）
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 页码字段
            run1 = fp.add_run("- ")
            self._apply_run_font(run1)

            fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run2 = fp.add_run()
            run2._r.append(fld_char_begin)

            instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run3 = fp.add_run()
            run3._r.append(instr)

            fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run4 = fp.add_run()
            run4._r.append(fld_char_end)

            run5 = fp.add_run(" -")
            run5.font.size = Pt(9)
            run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ═══════════════════════════════════════════════
    # PDF 导出
    # ═══════════════════════════════════════════════

    def to_pdf(
        self,
        docx_path: str,
        output_path: str,
        markdown_text: str = "",
    ) -> str:
        """
        将 .docx 转换为 PDF。

        优先使用 LibreOffice (subprocess)，fallback 到 reportlab。

        Args:
            docx_path: 输入的 .docx 文件路径
            output_path: 输出的 .pdf 文件路径

        Returns:
            str: 生成的 PDF 文件路径
        """
        if not output_path.endswith(".pdf"):
            output_path += ".pdf"

        # 方法1: LibreOffice
        if self._try_libreoffice_convert(docx_path, output_path):
            return output_path

        # 方法2: reportlab
        logger.info("LibreOffice 不可用，使用 reportlab 生成 PDF...")
        if not markdown_text:
            logger.warning("reportlab PDF 需要 markdown_text 参数")
            return output_path
        return self._reportlab_pdf(markdown_text, output_path)

    def _try_libreoffice_convert(self, docx_path: str, pdf_path: str) -> bool:
        """尝试用 LibreOffice 转换。"""
        try:
            import shutil

            soffice = shutil.which("libreoffice") or shutil.which("soffice")
            # P1-4（2026-08-07）：LibreOffice 安装时可能不在 PATH。
            # 追加常见安装路径的逐级探测，优先级从高到低。
            if not soffice:
                _candidates = [
                    # Windows
                    os.path.expandvars(r"%ProgramFiles%\LibreOffice\program\soffice.exe"),
                    os.path.expandvars(r"%ProgramFiles(x86)%\LibreOffice\program\soffice.exe"),
                    # macOS
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                    # Linux
                    "/usr/bin/soffice",
                    "/usr/local/bin/soffice",
                    "/opt/libreoffice/program/soffice",
                ]
                for cand in _candidates:
                    if os.path.isfile(cand):
                        soffice = cand
                        break
            if not soffice:
                logger.warning("LibreOffice 未安装或未在 PATH/常见路径中找到")
                return False

            out_dir = str(Path(pdf_path).parent)
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                logger.info(f"PDF (LibreOffice) 已保存: {pdf_path}")
                return True
            else:
                logger.warning(f"LibreOffice 转换失败: {result.stderr[:200]}")
                return False
        except Exception as e:
            logger.warning(f"LibreOffice 异常: {e}")
            return False

    def _reportlab_pdf(self, markdown_text: str, pdf_path: str) -> str:
        """用 reportlab 从 Markdown 生成 PDF（基础版本）。"""
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        # 注册中文字体
        font_registered = False
        for font_path_candidate in [
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]:
            if os.path.exists(font_path_candidate):
                try:
                    pdfmetrics.registerFont(TTFont("CJKFont", font_path_candidate))
                    font_registered = True
                    break
                except Exception:
                    continue

        if not font_registered:
            # Check for any available CJK font
            import glob

            for fp in glob.glob("/usr/share/fonts/**/*.ttf", recursive=True):
                if "droid" in fp.lower() or "noto" in fp.lower() or "cjk" in fp.lower():
                    try:
                        pdfmetrics.registerFont(TTFont("CJKFont", fp))
                        font_registered = True
                        break
                    except Exception:
                        continue

        font_name = "CJKFont" if font_registered else "Helvetica"

        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CJKTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            textColor=HexColor("#003366"),
            spaceAfter=12,
        )
        h1_style = ParagraphStyle(
            "CJKH1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=16,
            leading=20,
            textColor=HexColor("#003366"),
            spaceBefore=18,
            spaceAfter=8,
        )
        h2_style = ParagraphStyle(
            "CJKH2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            textColor=HexColor("#003366"),
            spaceBefore=14,
            spaceAfter=6,
        )
        h3_style = ParagraphStyle(
            "CJKH3",
            parent=styles["Heading3"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            textColor=HexColor("#003366"),
            spaceBefore=10,
            spaceAfter=5,
        )
        body_style = ParagraphStyle(
            "CJKBody",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=16,
            spaceAfter=6,
        )
        code_style = ParagraphStyle(
            "Code",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=8,
            leading=10,
            leftIndent=10,
        )

        story = []
        in_code_block = False

        for line in markdown_text.split("\n"):
            stripped = line.strip()

            if stripped.startswith("```"):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                story.append(Paragraph(stripped, code_style))
                continue

            if not stripped:
                story.append(Spacer(1, 6))
                continue

            # 标题
            hm = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if hm:
                level = len(hm.group(1))
                text = re.sub(r"\*\*(.*?)\*\*", r"\1", hm.group(2))
                style_map = {1: h1_style, 2: h2_style, 3: h3_style, 4: body_style}
                story.append(Paragraph(text, style_map.get(level, body_style)))
                continue

            # 表格
            if self._is_table_row(stripped):
                cells = self._parse_table_row(stripped)
                story.append(Paragraph(" | ".join(cells), body_style))
                continue

            # 引用
            if stripped.startswith(">"):
                text = stripped.lstrip("> ").strip()
                story.append(Paragraph(f"<i>{text}</i>", body_style))
                continue

            # 普通段落
            text = stripped
            text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
            text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
            story.append(Paragraph(text, body_style))

        doc.build(story)
        logger.info(f"PDF (reportlab) 已保存: {pdf_path}")
        return pdf_path

    # ═══════════════════════════════════════════════
    # 一键导出所有格式
    # ═══════════════════════════════════════════════

    def export_all(
        self,
        markdown_text: str,
        base_name: str,
        chart_paths: dict[str, str] | None = None,
        output_dir: str = "outputs",
        formats: list[str] | None = None,
    ) -> dict[str, str]:
        """
        一键导出所有格式。

        Args:
            markdown_text: Markdown 报告文本
            base_name: 输出文件名（不含扩展名）
            chart_paths: 图表路径字典
            output_dir: 输出目录
            formats: 需要的格式，默认 ["docx"]，可加 "pdf"

        Returns:
            dict: {"docx": "path", "pdf": "path"}
        """
        if formats is None:
            formats = ["docx"]

        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)

        results = {}

        if "docx" in formats:
            docx_path = str(out_dir / f"{base_name}.docx")
            results["docx"] = self.to_docx(markdown_text, docx_path, chart_paths)

        if "pdf" in formats:
            docx_path = results.get("docx", "")
            if not docx_path:
                docx_path = str(out_dir / f"{base_name}.docx")
                self.to_docx(markdown_text, docx_path, chart_paths)
            pdf_path = str(out_dir / f"{base_name}.pdf")
            results["pdf"] = self.to_pdf(docx_path, pdf_path, markdown_text)

        return results

    # ═══════════════════════════════════════════════════════
    def _add_cover_page(self, title_text="", date_str="", style_name=None):
        if style_name is None:
            style_name = self.style_id.upper() if self.style_id else "CICC"
        """添加专业封面页"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        body = self._doc.element.body
        existing = list(body)
        # 仅保留节属性元素（sectPr），丢弃模板残留段落（避免封面后出现大量空段/空 Heading）
        kept = []
        for child in existing:
            if child.tag.endswith("}sectPr"):
                kept.append(child)
            else:
                body.remove(child)
        existing = kept
        for _ in range(3):
            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(style_name)
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        r.font.name = "SimHei"
        r.bold = True
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u6df1\u5ea6\u7814\u7a76\u62a5\u544a")
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        r.font.name = "SimHei"
        r.bold = True
        for _ in range(1):
            self._doc.add_paragraph()
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u5206\u6790\u5e08\uff1a\u5468\u529b")
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        from datetime import datetime

        ds = date_str or datetime.now().strftime("%Y\u5e74%m\u6708%d\u65e5")
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ds)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        for _ in range(3):
            self._doc.add_paragraph()
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "\u672c\u62a5\u544a\u57fa\u4e8e\u516c\u5f00\u4fe1\u606f\uff0c\u4e0d\u6784\u6210\u6295\u8d44\u5efa\u8bae"
        )
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        self._doc.add_page_break()
        for child in existing:
            body.append(child)

    def _add_toc(self):
        """添加目录页"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u76ee  \u5f55")
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        r.font.name = "SimHei"
        self._doc.add_paragraph()
        toc_entries = []
        for para in self._doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading 1"):
                toc_entries.append(("", text))
            elif para.style.name.startswith("Heading 2"):
                toc_entries.append(("  ", text))
        for indent, text in toc_entries:
            p = self._doc.add_paragraph()
            r = p.add_run(indent + text)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r.font.name = "SimSun"
            p.paragraph_format.space_after = Pt(2)
        self._doc.add_page_break()


# CLI 入口
# ═══════════════════════════════════════════════════════


def main():
    import argparse

    parser = argparse.ArgumentParser(description="2号分析师 — 报告导出器")
    parser.add_argument("--input", "-i", default="test_report.md", help="输入 Markdown 文件")
    parser.add_argument("--output", "-o", default="outputs/test_report", help="输出文件名（不含扩展名）")
    parser.add_argument("--company", "-c", default="测试公司", help="公司名称（用于页眉）")
    parser.add_argument("--pdf", action="store_true", help="同时导出 PDF")
    parser.add_argument("--charts", "-g", default=None, help="图表目录（自动查找 PNG）")

    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(message)s")

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误: 输入文件不存在: {args.input}")
        sys.exit(1)

    text = input_path.read_text(encoding="utf-8")

    chart_paths = {}
    if args.charts:
        charts_dir = Path(args.charts)
        if charts_dir.is_dir():
            for png in sorted(charts_dir.glob("*.png")):
                chart_paths[png.stem] = str(png)
            print(f"找到 {len(chart_paths)} 个图表文件")

    exporter = ReportExporter(company_name=args.company)

    formats = ["docx"]
    if args.pdf:
        formats.append("pdf")

    result = exporter.export_all(
        markdown_text=text,
        base_name=args.output.rstrip(".docx").rstrip(".pdf"),
        chart_paths=chart_paths,
        formats=formats,
    )

    print(f"\n{'=' * 60}")
    print("导出完成!")
    for fmt, path in result.items():
        size = Path(path).stat().st_size
        print(f"  {fmt}: {path} ({size / 1024:.1f} KB)")
    print(f"{'=' * 60}")
